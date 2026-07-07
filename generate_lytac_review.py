#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate LYTAC review document with specified formatting."""

import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

# 参考文献编号与英文题名（正文标注用，不计入正文字数）
REF_TITLES = {
    1: "Lysosome-targeting chimaeras for degradation of extracellular proteins",
    2: "LYTACs that engage the asialoglycoprotein receptor for targeted protein degradation",
    3: "Modular cytokine receptor-targeting chimeras for targeted degradation of cell surface and extracellular proteins",
    4: "Covalent LYTAC enabled by DNA aptamers for immune checkpoint degradation therapy",
    5: "Biomarker-activated multifunctional lysosome-targeting chimeras mediated selective degradation of extracellular amyloid fibrils",
    6: "Aptamer-LYTACs for targeted degradation of extracellular and membrane proteins",
    7: "Targeted degradation of membrane and extracellular proteins with LYTACs",
    8: "Targeted degradation of extracellular secreted and membrane proteins",
    9: "Bifunctional small molecules that mediate the degradation of extracellular proteins",
    10: "PROTAC targeted protein degraders: the past is prologue",
    11: "Emerging protein degradation strategies: expanding the scope to extracellular and membrane proteins",
    12: "Designed endocytosis-inducing proteins degrade targets and amplify signals",
    13: "Insulin-like Growth Factor 2 (IGF2)-Fused Lysosomal Targeting Chimeras for Degradation of Extracellular and Membrane Proteins",
    14: "Lysosome-Targeting Chimera Using Mannose-6-Phosphate Glycans Derived from Glyco-Engineered Yeast",
    15: "Lysosome-targeting chimeras containing an endocytic signaling motif trigger endocytosis and lysosomal degradation of cell-surface proteins",
    16: "Targeted degradation of extracellular proteins: state of the art and diversity of degrader designs",
    17: "Targeted protein degradation: advances in drug discovery and clinical practice",
    18: "Targeted protein degradation: mechanisms, strategies and application",
}


def cite(*nums):
    """在段落文本中暂存文献编号；实际显示时集中移到段末。"""
    return "".join(f"{{CITE:{n}}}" for n in nums)


def format_citations(text):
    """将分散在句中的引用集中为段末文献名标注，使正文阅读更连贯。"""
    refs = []

    def collect(match):
        ref_id = int(match.group(1))
        if ref_id not in refs:
            refs.append(ref_id)
        return ""

    body = re.sub(r"\{CITE:(\d+)\}", collect, text)
    if not refs:
        return body
    note = "（参考文献：" + "；".join(f"《{REF_TITLES[n]}》[{n}]" for n in refs) + "）"
    return body + note


def count_chinese_chars(text):
    """统计汉字数；排除文献名标注块"""
    cleaned = re.sub(r"\{CITE:\d+\}", "", text)
    cleaned = re.sub(r"（参考文献：[^）]*）", "", cleaned)
    return sum(1 for c in cleaned if "\u4e00" <= c <= "\u9fff")


def set_run_font(run, chinese_font="仿宋", western_font="Times New Roman", size_pt=12, bold=False):
    run.font.name = western_font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), chinese_font)


def add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=True, bold=False, size_pt=12):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    if first_line_indent:
        pf.first_line_indent = Cm(0.74)
    run = p.add_run(format_citations(text))
    set_run_font(run, size_pt=size_pt, bold=bold)
    return p


def add_heading_text(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, size_pt=sizes.get(level, 12), bold=True)
    return p


def add_figure_placeholder(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("【此处插入示意图】")
    set_run_font(run, size_pt=12)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, size_pt=10.5)


def add_reference(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0.74)
    pf.hanging_indent = Cm(0.74)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size_pt=10.5)


ABSTRACT = (
    "靶向蛋白降解是近年来药物研发领域备受关注的研究方向之一。"
    "与传统抑制或阻断靶蛋白功能不同，该类策略旨在将致病蛋白导向细胞内降解系统。"
    "溶酶体靶向嵌合体（Lysosome-Targeting Chimera, LYTAC）是其中新兴的技术类型，"
    "主要利用细胞表面溶酶体靶向受体，将胞外蛋白及膜蛋白转运至溶酶体降解。"
    "鉴于PROTAC主要作用于胞内蛋白，LYTAC为胞外靶点的干预提供了补充性技术路径。"
    "本综述在梳理相关文献的基础上，介绍LYTAC的基本作用方式，"
    "并结合Banik等提出的开创性工作"
    + cite(1)
    + "、Ahn等实现的肝组织特异性降解"
    + cite(2)
    + "以及Pance等报道的KineTAC蛋白平台"
    + cite(3)
    + "，讨论该技术的发展过程与可能的应用方向。"
    "现有研究表明，LYTAC能否真正走向临床，不仅取决于分子设计，"
    "还取决于对内吞、受体循环和溶酶体分选等生物学过程的认识程度。"
    "目前该领域仍以基础研究为主，临床证据仍然较少。"
)

INTRO = [
    "多种疾病均与蛋白质表达异常或功能失调相关，如肿瘤、自身免疫病和神经退行性疾病。"
    "传统药物研发往往通过小分子或抗体去结合靶蛋白的活性位点或结合位点，从而抑制其功能。"
    "但对于一些没有明确药物结合口袋的蛋白，或者像自身抗体这类靶点，单纯阻断往往效果有限，"
    "因此常被称为较难成药的靶点。",
    "近二十年来，靶向蛋白降解技术发展较快。PROTAC和分子胶等策略已经能把部分胞内蛋白"
    "导向泛素-蛋白酶体系统降解，其中一些分子已进入临床试验。"
    "Bekes等在综述中回顾了PROTAC的发展历程"
    + cite(10)
    + "，并指出蛋白降解药物正在从概念验证走向临床验证。"
    "Zhong等的综述进一步从临床转化角度梳理了TPD领域的进展与挑战"
    + cite(17)
    + "。",
    "不过PROTAC通常要求靶蛋白具有可被配体结合的胞内区域，这样才能招募E3连接酶。"
    "而分泌至细胞外的蛋白及跨膜蛋白数量众多，在多种疾病中亦十分常见，"
    "却长期未纳入传统PROTAC的主要作用范围。",
    "从生物医学角度看，这类蛋白参与细胞通讯、免疫识别和信号转导。"
    "例如PD-L1、EGFR、IgE以及多种自身抗体相关复合物，都属于胞外或膜蛋白。"
    "它们有时表达量高，有时更新较快，单纯阻断未必能长期控制疾病进程。"
    "因此，能否把这些蛋白送入溶酶体彻底降解，是一个既有科学问题也有应用价值的问题。",
    "2020年，Banik等在Nature上提出了溶酶体靶向嵌合体（LYTAC）的概念"
    + cite(1)
    + "。"
    "此后几年里，不同研究组在分子设计、受体选择、细胞机制和疾病模型等方面开展了大量研究。",
    "Li等和Chen等发表的综述文章，对LYTAC技术的发展脉络和应用前景做了较系统的介绍"
    + cite(7, 8)
    + "。"
    "Lin等的综述从更广的胞外蛋白降解角度，将LYTAC与AbTAC等策略置于同一框架下讨论"
    + cite(11)
    + "。"
    "Mamun等的综述则系统归纳了胞外靶向蛋白降解（eTPD）的技术类型与成药性挑战"
    + cite(16)
    + "。",
    "本综述在系统梳理上述文献的基础上，"
    "归纳LYTAC技术的作用机制、代表性研究进展及应用前景，"
    "并讨论生物医学基础认识对这类新药研发思路的影响。",
]

BODY_SECTION1 = [
    "LYTAC本质上是一种双功能分子，一端结合靶蛋白，另一端结合细胞表面的溶酶体靶向受体"
    "（Lysosome-Targeting Receptor, LTR），中间通过连接子相连。"
    "靶蛋白结合部分可以用小分子、肽、抗体或适体，受体结合部分则常用甘露糖-6-磷酸聚糖或GalNAc等配体。"
    "当LYTAC同时结合靶蛋白和受体后，会形成三元复合物，随后通过内吞进入内体，"
    "再进一步分选到溶酶体，靶蛋白最终被水解（图1）。"
    "Chen等在综述中指出"
    + cite(8)
    + "，LYTAC属于依赖溶酶体的胞外蛋白降解策略，与PROTAC依赖蛋白酶体的路径明显不同。"
    "Li等在综述中把LYTAC列为靶向蛋白降解领域较新的技术方向之一"
    + cite(7)
    + "，并进一步说明胞外蛋白和膜蛋白在肿瘤、自身免疫病和神经退行性疾病中都很常见。"
    "Lin等的综述指出，胞外和膜蛋白约占蛋白质组总量的相当比例，"
    "在多种疾病发病机制中发挥关键作用"
    + cite(11)
    + "。",
    "Mamun等在综述中将LYTAC归入胞外靶向蛋白降解（eTPD）的代表性技术，"
    "认为其通过把靶蛋白与细胞表面LTR连接，经内吞-溶酶体途径实现蛋白清除"
    + cite(16)
    + "。",
    "这个过程和抗体单纯阻断靶蛋白并不相同。抗体通常只是占据结合位点，"
    "而LYTAC的目标是减少靶蛋白本身的含量。"
    "文献中一般认为，这类分子属于事件驱动型作用方式，"
    "即结合剂无需具备抑制活性，只要能有效识别并捕获靶蛋白即可。"
    "由于降解发生在溶酶体，靶蛋白的信号传导功能和蛋白相互作用都可能随之减弱，"
    "这一点在EGFR和PD-L1等靶点上已有实验支持"
    + cite(1, 4)
    + "。",
    "与单纯阻断相比，直接清除致病蛋白有望更彻底地干预疾病进程，"
    "但实际疗效仍受降解效率、靶蛋白更新动力学及组织分布等因素制约。",
    "目前已报道的LTR包括CI-M6PR（也称IGF2R）、ASGPR、整合素以及部分细胞因子受体等。"
    "Chen等在综述中归纳了多种LTR及其在胞外蛋白降解中的应用"
    + cite(8)
    + "，Lin等的综述亦对LYTAC依赖的受体类型做了系统梳理"
    + cite(11)
    + "。",
    "不同受体在组织分布和内吞特性上差别较大，受体选择会直接影响LYTAC作用的位置和效率。"
    "CI-M6PR在多种细胞上都有表达，因此第一代LYTAC的作用范围较广，"
    "但也更容易带来脱靶降解的风险。ASGPR主要在肝细胞表达，"
    "因此GalNAc-LYTAC更适合肝相关靶点"
    + cite(2, 6)
    + "。",
    "受体可视为药物进入细胞的分子入口，受体类型的选择实质上决定了药物降解作用的主要组织部位。"
    "LYTAC疗效的解读不仅依赖分子结构，"
    "还取决于受体占用、复合物回收及内体向溶酶体分选等细胞生物学过程。"
    "分子设计与细胞生物学机制的有机结合，是理解相关实验结果的重要前提。",
    "在设计分子时，需要同时考虑靶蛋白结合剂、受体配体和连接子。"
    "Banik等最初的LYTAC使用了抗体或小分子作为靶向模块"
    + cite(1)
    + "，表明该技术可根据靶蛋白特点选择不同的结合方式。"
    "受体配体的选择和价态同样关键。第一代LYTAC多使用CI-M6PR配体，"
    "通常采用多价甘露糖-6-磷酸聚糖以提高结合能力。"
    "Ahn等后来改用肝特异性受体ASGPR的tri-GalNAc配体"
    + cite(2)
    + "，说明组织分布是决定降解位置的重要因素。",
    "Caianiello等报道了小分子降解胞外蛋白的MoDE-A策略"
    + cite(9)
    + "，通过ASGPR把靶蛋白送入溶酶体，"
    "说明除抗体-聚糖偶联物外，小分子也可介导胞外蛋白降解。"
    "Zhang等进一步将IGF2与抗体或纳米抗体融合，构建了基因编码的iLYTAC"
    + cite(13)
    + "，"
    "可在细菌表达系统中大量制备，并降解EGFR、PD-L1及α-突触核蛋白等靶点。",
    "Kim等则利用糖工程酵母来源的M6P聚糖制备LYTACgyM6pG"
    + cite(14)
    + "，"
    "以无铜点击化学将聚糖与抗PD-L1纳米抗体偶联，"
    "在简化聚糖配体制备的同时实现了PD-L1的溶酶体降解。",
    "Li等在总结影响降解效率的因素时"
    + cite(7)
    + "，"
    "提到了配体价态、连接子长度和结合表位等分子层面的变量，"
    "并强调分子设计和细胞环境需要同时考虑。",
    "Bekes等在回顾PROTAC发展时指出"
    + cite(10)
    + "，"
    "该领域已从学术研究走向产业开发，并有PROTAC分子进入临床前和早期临床试验阶段。"
    "这一发展轨迹对LYTAC的未来路径具有参照意义。",
    "结合Chen等的综述"
    + cite(8)
    + "，可将LYTAC与PROTAC进行如下对比。"
    "PROTAC主要连接靶蛋白和E3连接酶，把蛋白送进蛋白酶体降解，更适合有胞内结构域的靶点。"
    "LYTAC则连接靶蛋白和溶酶体靶向受体，把蛋白送进溶酶体降解，"
    "更适合分泌蛋白、抗体和膜蛋白。",
    "这两种技术并非相互替代关系，而是针对不同蛋白亚细胞定位的互补策略。",
    "从细胞生物学角度看，受体介导的内吞并不是单向过程。"
    "复合物进入早期内体后，一部分会被分选到溶酶体，"
    "另一部分则可能通过retromer等回收机器回到细胞膜。"
    "Lin等的综述和Mamun等的综述均指出，"
    "受体内吞后的分选与回收是影响降解效率的重要环节"
    + cite(11, 16)
    + "。",
    "Huang等设计的EndoTag蛋白可与IGF2R、ASGPR等受体结合并触发内吞"
    + cite(12)
    + "，"
    "其研究亦提示，内源性配体与外源降解分子之间可能存在受体结合竞争，"
    "这一因素在解读不同细胞系中的降解效率差异时值得关注。",
    "Fang等开发了含CI-M6PR来源内吞信号肽的SignalTAC"
    + cite(15)
    + "，"
    "在不依赖外源LTR配体的条件下即可促进靶蛋白内吞和溶酶体降解，"
    "为分子设计提供了另一条思路。",
    "在分子设计层面，多价配体虽然能提高与受体的结合能力，"
    "但也可能带来非特异性聚集或过快清除。"
    "Li等在综述中指出"
    + cite(7)
    + "，"
    "连接子过长可能导致两个功能模块之间距离过大，不利于三元复合物形成，"
    "而过短又可能限制构象调整。",
    "LYTAC研究要求将化学结构设计与膜受体、内吞小泡及溶酶体等细胞过程相联系。"
    "例如溶酶体抑制剂实验、受体敲除实验及遗传学筛选，"
    "都是用来证明降解确实走了溶酶体途径的对照手段"
    + cite(1)
    + "。",
    "从靶点选择角度看，膜蛋白和分泌蛋白在疾病中的角色各不相同。"
    "有些蛋白主要起信号传导作用，如EGFR和HER2；"
    "有些则属于免疫抑制分子或致病聚集物，如PD-L1和Aβ。"
    "因此，同一个LYTAC技术平台面对不同疾病时，评价指标亦不尽相同，"
    "不宜仅以体外降解百分比作为唯一指标。",
    "Mamun等在综述中指出"
    + cite(16)
    + "，"
    "胞外靶向蛋白降解的核心逻辑是把靶蛋白送入溶酶体，"
    "而不是像PROTAC那样依赖泛素-蛋白酶体系统。"
    "Zhao等在靶向蛋白降解综述中亦提到，"
    "溶酶体途径是拓展降解靶点范围的重要方向之一"
    + cite(18)
    + "。",
    "在结合模块方面，Wu等和Li等分别报道了适体连接的Apt-LYTAC和共价适体LYTAC"
    + cite(6, 4)
    + "，"
    "证明核酸适体可替代抗体作为靶蛋白识别模块，且合成与筛选相对方便。",
    "Huang等和Pance等则从不同角度发展了基因编码的蛋白降解平台"
    + cite(12, 3)
    + "，"
    "为绕开复杂化学偶联提供了可行路径。",
    "Banik等的研究揭示了CI-M6PR的生理学背景"
    + cite(1)
    + "："
    "该受体负责介导带有M6P修饰的溶酶体酶进入细胞，"
    "LYTAC正是利用了这一受体介导的货物分拣通路。",
    "Ahn等在GalNAc-LYTAC研究中发现，在抗体Fc区进行定点偶联可改善体内药代动力学表现"
    + cite(2)
    + "，"
    "说明连接位点和分子构象对体内行为具有重要影响。",
    "Zhao等在靶向蛋白降解综述中系统比较了蛋白酶体途径与溶酶体途径降解策略的适用范围，"
    "指出LYTAC、AbTAC及适体嵌合体等属于溶酶体依赖的胞外降解技术"
    + cite(18)
    + "。",
    "从制备策略看，化学合成的多价M6P聚糖配体合成步骤繁琐，"
    "且批次间均一性难以保证；"
    "Kim等利用糖工程酵母发酵制备gyM6pG聚糖，"
    "再以无铜点击化学与纳米抗体偶联，"
    "为LYTAC的规模化生产提供了新思路"
    + cite(14)
    + "。",
    "Zhang等报道的iLYTAC采用IGF2与靶蛋白结合域的基因融合表达策略，"
    "可在数天内通过标准克隆和细菌表达获得目标蛋白，"
    "并实现对EGFR、PD-L1、CD20等靶点的高效降解"
    + cite(13)
    + "。",
    "Huang等设计的EndoTag通过计算蛋白质设计获得，"
    "可与IGF2R、ASGPR、转铁蛋白受体等多种LTR结合，"
    "融合靶蛋白结合蛋白后形成pLYTAC，"
    "在模块化、遗传可编码性及组织靶向性方面展现出优势"
    + cite(12)
    + "。",
    "Fang等的SignalTAC将CI-M6PR来源的10个氨基酸酪氨酸基内吞信号肽P3"
    "与靶蛋白结合模块融合，"
    "在不依赖外源LTR配体的情况下即可诱导靶蛋白内吞和溶酶体降解，"
    "且抗肿瘤活性优于亲本抗体"
    + cite(15)
    + "。",
    "Ahn等在GalNAc-LYTAC研究中除EGFR外，"
    "还以肽段为结合剂实现了整合素降解，"
    "并抑制了癌细胞增殖，"
    "说明较小的结合模块亦可有效介导溶酶体降解"
    + cite(2)
    + "。",
    "Wu等的Apt-LYTAC研究采用了类似的肝靶向策略，"
    "但用适体替代抗体作为靶蛋白结合模块，"
    "在简化分子构建的同时实现了PDGF和PTK7的降解"
    + cite(6)
    + "。",
    "上述研究表明，LYTAC分子设计正在从单一的聚糖-抗体偶联模式，"
    "向小分子、适体、融合蛋白、内吞信号肽及病灶激活前体等多元化方向发展。",
    "在内吞与分选机制方面，Banik等通过CRISPR干扰筛选发现，"
    "exocyst复合体参与CI-M6PR在细胞膜的呈现，"
    "提示LYTAC-受体复合物的胞内转运受多种细胞机器调控"
    + cite(1)
    + "。"
    "Chen等在综述中进一步指出，"
    "整合素、清道夫受体及细胞因子受体等亦可作为组织选择性降解的效应器，"
    "拓展了LTR的选择范围"
    + cite(8)
    + "。",
    "Caianiello等开发的MoDE-A属于双功能小分子降解剂，"
    "一端结合ASGPR，另一端结合靶蛋白，"
    "在肝细胞中实现了整合素等靶点的溶酶体降解"
    + cite(9)
    + "。"
    "与抗体-聚糖偶联物相比，小分子降解剂分子量更低、"
    "结构表征更简单，但靶点结合亲和力及体内稳定性仍需优化。",
    "从事件驱动药理角度看，LYTAC不要求结合剂具备酶抑制活性，"
    "这一特点使其能够靶向传统抑制剂难以作用的蛋白界面。"
    "Li等在综述中将其与抗体阻断策略进行了比较，"
    "认为直接降解靶蛋白有望更彻底地消除其信号传导和蛋白相互作用功能"
    + cite(7)
    + "。",
    "然而，降解效率还受到靶蛋白合成与更新速率的影响。"
    "对于更新较快的膜蛋白，"
    "即使实现了高效的溶酶体清除，"
    "细胞仍可能通过持续合成维持一定水平的蛋白表达，"
    "从而削弱降解剂的长期药效。",
    "因此，在评价LYTAC疗效时，"
    "除监测蛋白水平变化外，"
    "还应关注下游信号通路抑制程度及细胞功能表型改变。",
    "Zhong等的综述指出，"
    "TPD药物研发需兼顾靶点可降解性、分子理化性质及联合治疗策略等因素，"
    "这些原则对LYTAC同样适用"
    + cite(17)
    + "。",
    "Bekes等回顾PROTAC时提到，"
    "随着更多E3连接酶和靶点类型被纳入降解设计，"
    "该领域正从肿瘤适应症向更广泛的疾病领域扩展"
    + cite(10)
    + "。"
    "LYTAC由于可作用于分泌蛋白和膜蛋白，"
    "在自身免疫病、代谢病及神经退行性疾病中均具有潜在应用空间，"
    "但相关研究尚处于早期阶段。",
    "Lin等的综述对LYTAC的作用机制进行了图示化归纳："
    "LYTAC同时结合POI和LTR后形成三元复合物，"
    "经网格蛋白介导的内吞进入内体，"
    "再经分选进入溶酶体完成蛋白水解，"
    "而LTR则可经回收途径返回细胞膜"
    + cite(11)
    + "。"
    "理解这一循环过程，有助于解释为何同一LYTAC在不同细胞系中可能呈现显著不同的降解效率。",
]

BODY_SECTION2 = [
    "2020年，Banik等在Nature发表了LYTAC领域的开创性工作"
    + cite(1)
    + "。"
    "作者将靶向模块与CI-M6PR配体通过连接子连接，构建了首批LYTAC分子。"
    "他们还利用CRISPR筛选研究了LYTAC的内吞通路，发现exocyst复合体参与CI-M6PR的膜面呈现，"
    "表明该研究不仅涉及分子构建，也在补充受体介导内吞的基础认识。"
    "实验中，LYTAC成功降解了ApoE4、EGFR、CD71和PD-L1等蛋白。"
    "以EGFR为例，处理后受体水平在数小时内明显下降，溶酶体抑制剂可以阻断这一效应，"
    "说明溶酶体途径是主要降解路径。PD-L1被降解后，肿瘤细胞的免疫逃逸能力也有所减弱（图2）。",
    "第一代LYTAC主要依赖CI-M6PR，而该受体在多种组织中都有表达，组织选择性不够理想。"
    "2021年，Ahn等在Nature Chemical Biology报道了结合ASGPR的GalNAc-LYTAC"
    + cite(2)
    + "。"
    "研究者将靶蛋白结合剂与tri-GalNAc配体连接，在肝细胞中实现了EGFR降解。"
    "结果显示，与单纯使用抑制性抗体相比，GalNAc-LYTAC对EGFR下游信号的抑制更持久。"
    "作者还把较小的肽段结合剂与tri-GalNAc连接，实现了整合素降解并抑制癌细胞增殖。"
    "此外，他们在抗体Fc区进行定点偶联，改善了GalNAc-LYTAC在体内的药代动力学表现（图3）。",
    "Li等把这类工作归纳为第二代组织特异性LYTAC"
    + cite(7)
    + "，"
    "认为肝靶向设计对减少脱靶效应具有一定帮助。",
    "化学合成LYTAC在制备和均一性方面有一定难度，"
    "因此也有研究尝试用基因编码的蛋白分子来实现类似功能。",
    "2023年，Pance等在Nature Biotechnology提出了KineTAC平台"
    + cite(3)
    + "。"
    "KineTAC是一种双特异性抗体，一条臂结合细胞因子受体，另一条臂结合靶蛋白。"
    "研究者针对PD-L1、HER2、PD-1、EGFR、CDCP1和TROP2等靶点构建了KineTAC，"
    "最大降解效率大约在51%到93%之间。"
    "与化学LYTAC相比，KineTAC可通过哺乳动物细胞表达生产，"
    "更换靶蛋白结合臂也相对方便（图4）。",
    "Chen等把KineTAC与化学合成的LYTAC并列讨论"
    + cite(8)
    + "，"
    "认为这类蛋白平台在制备和替换靶点结合臂方面可能更方便。",
    "Wu等在Angewandte Chemie International Edition报道了Apt-LYTAC"
    + cite(6)
    + "，"
    "把适体与tri-GalNAc连接，实现了肝细胞中对PDGF和PTK7等靶点的降解。",
    "Li等在Journal of the American Chemical Society报道了基于DNA适体的共价LYTAC"
    + cite(4)
    + "，"
    "可在体内降解PD-L1，并观察到比传统免疫检查点抗体更强的抗肿瘤免疫反应。",
    "Zhang等报道的iLYTAC通过IGF2融合蛋白实现溶酶体靶向"
    + cite(13)
    + "，"
    "并在异种移植瘤模型中验证了抗肿瘤活性。",
    "Huang等设计的EndoTag融合蛋白（pLYTAC）"
    + cite(12)
    + "，"
    "在小鼠肿瘤模型中显示出优于单纯抗体的PD-L1靶向降解效果。",
    "这些研究共同说明，LYTAC已经从最初的概念验证，"
    "逐步发展到组织靶向、适体平台、蛋白工程平台及聚糖制备工艺优化等多个方向。",
    "综合上述研究可见，该领域的发展并非单一线性推进，"
    "而是同时在分子类型、受体选择和制备策略三个层面不断扩展。",
    "Banik等的Nature论文的重要意义不仅在于提出了LYTAC这一概念，"
    "更在于首次较完整地证明该溶酶体降解路线可推广至多个疾病相关靶点"
    + cite(1)
    + "。",
    "Ahn等的工作实现了组织特异性降解"
    + cite(2)
    + "，"
    "Pance等的KineTAC则提供了另一条实现路径"
    + cite(3)
    + "，"
    "说明胞外蛋白降解并不必然依赖化学合成的聚糖-蛋白偶联物。",
    "Wu等和Li等的适体研究表明结合模块可进一步小型化、更灵活"
    + cite(6, 4)
    + "。",
    "进一步比较这几类平台，化学合成的聚糖-蛋白LYTAC在结构修饰上较灵活，"
    "可以精细调整配体价态和连接子，但制备难度也更高。"
    "Kim等利用酵母来源M6P聚糖的策略"
    + cite(14)
    + "，"
    "在一定程度上缓解了合成聚糖配体的复杂性。",
    "KineTAC和iLYTAC依托蛋白工程，在更换靶点结合臂时相对方便；"
    "适体LYTAC则分子量较小，在体外筛选不同结合序列时有一定优势。",
    "Liu等在Chem上发表了病灶激活LYTAC前体"
    + cite(5)
    + "，"
    "利用阿尔茨海默病病灶区铜离子催化点击反应，在局部生成活性分子，"
    "并通过CD206受体促进Aβ聚集体进入溶酶体降解。"
    "这类研究把疾病局部病理特征和药物激活结合起来，"
    "试图降低全身给药带来的脱靶风险。",
    "Fang等的SignalTAC研究"
    + cite(15)
    + "，"
    "则通过融合内吞信号肽增强膜蛋白向溶酶体的转运，"
    "在抗肿瘤实验中显示出优于亲本抗体的疗效。",
    "综合来看，LYTAC领域正从可行性验证阶段转向精准性、稳定性及可开发性优化阶段。",
    "Bertozzi团队更强调受体化学、细胞机制和分子平台的系统搭建"
    + cite(1, 2)
    + "，"
    "国内和亚洲其他实验室则在适体偶联、病灶激活、IGF2融合蛋白及酵母聚糖制备等方面开展了系列拓展研究"
    + cite(4, 5, 6, 13, 14)
    + "。",
    "如果把时间线拉长来看，PROTAC的快速发展为蛋白降解药物提供了重要参照。"
    "Bekes等回顾这一历程时指出"
    + cite(10)
    + "，"
    "PROTAC在取得临床概念验证后，产业界投入明显增加。"
    "Zhong等的综述从临床数据角度进一步说明，"
    "TPD技术正在从实验室走向临床试验"
    + cite(17)
    + "。",
    "LYTAC目前处于类似但更早期的阶段："
    "已有高质量概念验证论文，也有多个并行平台，但尚缺少系统的临床数据。",
    "Banik等用CRISPRi筛选内吞相关基因"
    + cite(1)
    + "，"
    "体现了LYTAC研究对细胞生物学机制的持续关注。",
    "上述研究模式将药物化学与细胞生物学置于同等重要的地位，"
    "体现了机制导向药物研发的基本理念。",
    "Lin等的综述对LYTAC与AbTAC、KineTAC等技术进行了横向比较，"
    "指出不同eTPD平台在靶点类型、制备难度和体内行为方面各有特点"
    + cite(11)
    + "。",
    "Mamun等的综述进一步按分子类型将现有降解剂分为LYTAC、双特异性抗体、"
    "小分子降解剂及受体非依赖策略等类别，"
    "并归纳了肿瘤相关靶点的研究进展"
    + cite(16)
    + "。",
    "在适体平台方面，Wu等将识别靶蛋白的适体与tri-GalNAc连接，"
    "构建了可在肝细胞中降解PDGF和PTK7的Apt-LYTAC"
    + cite(6)
    + "。"
    "Li等则通过生物正交反应增强适体与PD-L1的共价结合，"
    "在肿瘤模型中实现了免疫检查点蛋白的降解及免疫激活"
    + cite(4)
    + "。",
    "两项研究共同表明，适体模块具有分子量小、易于化学修饰和快速筛选等优势，"
    "是抗体之外值得重视的靶蛋白识别元件。",
    "在蛋白工程平台方面，Pance等的KineTAC利用细胞因子受体介导的内吞，"
    "Huang等的EndoTag利用人工设计的内吞诱导蛋白，"
    "Zhang等的iLYTAC利用IGF2介导的CI-M6PR通路，"
    "三者均避免了复杂聚糖化学合成，"
    "但各自面临蛋白稳定性、免疫原性或受体竞争等不同挑战"
    + cite(3, 12, 13)
    + "。",
    "在神经退行性疾病方向，Liu等设计的病灶激活LYTAC前体"
    "利用阿尔茨海默病脑内Aβ沉积区高浓度铜离子触发点击反应，"
    "在局部原位生成活性降解分子，"
    "并通过CD206受体介导Aβ纤维的溶酶体清除"
    + cite(5)
    + "。"
    "该策略将疾病生物标志物纳入药物设计逻辑，"
    "为降低全身给药脱靶效应提供了范例。",
    "Zhong等的综述从PROTAC临床进展出发，"
    "指出TPD药物研发正经历从概念验证到临床转化的关键阶段，"
    "为评估LYTAC等新兴技术的成药前景提供了参照框架"
    + cite(17)
    + "。",
    "Kim等报道的LYTACgyM6pG在降解PD-L1后，"
    "可显著增强T细胞对肿瘤细胞的杀伤活性，"
    "其抗肿瘤效应优于单独使用抗PD-L1纳米抗体"
    + cite(14)
    + "。"
    "该研究不仅验证了酵母来源聚糖在LYTAC中的应用价值，"
    "也为免疫检查点蛋白的降解治疗提供了实验依据。",
    "Banik等在初代LYTAC研究中测试了ApoE4的降解，"
    "表明LYTAC并非仅面向肿瘤靶点，"
    "在神经退行性疾病相关分泌蛋白上亦具有探索潜力"
    + cite(1)
    + "。"
    "Liu等在此基础上进一步开发了针对Aβ纤维的病灶激活策略，"
    "将LYTAC的应用场景拓展至需要局部激活的复杂病理环境"
    + cite(5)
    + "。",
    "Pance等在KineTAC研究中测试了PD-1、TROP2等多个免疫和肿瘤相关靶点，"
    "最大降解效率介于51%至93%之间，"
    "表明细胞因子受体介导的内吞途径具有较广的靶点适用范围"
    + cite(3)
    + "。",
    "Huang等则证明，EndoTag与PD-L1抗体融合后，"
    "在小鼠肿瘤模型中的疗效显著优于抗体单药，"
    "且不同EndoTag因受体组织分布差异可实现降解的组织定向"
    + cite(12)
    + "。",
]

BODY_SECTION3 = [
    "从已发表研究来看，LYTAC已经在肿瘤、免疫相关疾病和神经退行性疾病等方向"
    "开展了较多探索，但多数仍停留在细胞和动物实验阶段。",
    "在肿瘤免疫方向，Banik等最早证明LYTAC可以降解PD-L1"
    + cite(1)
    + "，"
    "Li等进一步用共价适体LYTAC在动物模型中验证了抗肿瘤效果"
    + cite(4)
    + "，"
    "Pance等的KineTAC研究和Huang等的pLYTAC研究也证明PD-L1、HER2等靶点可以通过溶酶体途径被降解"
    + cite(3, 12)
    + "。",
    "PD-L1是免疫检查点治疗中非常重要的靶点，目前临床主要使用抗体阻断PD-1/PD-L1相互作用。"
    "Li等的工作说明，若将PD-L1直接降解，"
    "可能在某些肿瘤模型中产生比单纯阻断更强的免疫激活效果"
    + cite(4)
    + "。",
    "不过这类结果目前仍主要来自动物实验，尚不能直接外推至人体疗效。",
    "在肝靶向方向，Ahn等的GalNAc-LYTAC和Wu等的Apt-LYTAC都说明，"
    "选择合适的肝特异性受体是减少全身脱靶的重要思路"
    + cite(2, 6)
    + "。",
    "组织特异性受体在LYTAC研发中具有关键意义，"
    "是决定候选药物是否具备开发价值的重要因素之一。",
    "在神经退行性疾病方向，Liu等设计了可在阿尔茨海默病病灶区激活的LYTAC前体"
    + cite(5)
    + "，"
    "利用Aβ沉积区铜离子催化点击反应，在局部生成活性分子，"
    "并通过CD206受体促进Aβ聚集体进入溶酶体降解。"
    "作者还构建了可穿越血脑屏障的纳米载体，把点击反应前体送到脑部，"
    "再在病灶区原位生成活性降解分子。",
    "Li等认为，这类病灶激活设计对神经退行性疾病可能更有意义"
    + cite(7)
    + "，"
    "因为可以避免药物在全身过早发挥作用。",
    "Chen等和Li等都提到，分泌蛋白和膜蛋白是LYTAC较有潜力的靶点类型"
    + cite(8, 7)
    + "，"
    "但现有证据主要来自细胞和动物实验，真正进入临床阶段的项目仍然很少。",
    "Zhong等和Zhao等的综述均指出，"
    "LYTAC及相关溶酶体降解策略仍主要处于临床前阶段"
    + cite(17, 18)
    + "。",
    "虽然LYTAC研究进展较快，但要真正用于新药开发仍面临诸多挑战。",
    "含聚糖的LYTAC结构复杂，化学合成和生物偶联产物不够均一，给质量控制带来困难。"
    "Kim等的研究表明，利用糖工程酵母制备M6P聚糖偶联物"
    + cite(14)
    + "，"
    "有望改善聚糖配体的可及性与可放大性。",
    "CI-M6PR表达较广，可能导致非靶组织也发生蛋白降解。"
    "ASGPR等组织特异性受体有所改善，但其他器官的靶向仍缺少成熟方案。",
    "第一代LYTAC往往和靶蛋白一起被降解，体内有效浓度维持时间较短。",
    "Chen等在讨论胞外蛋白降解策略时指出"
    + cite(8)
    + "，"
    "连接子设计、受体占用和细胞内吞分选都会影响最终降解效果。"
    "Li等也总结了LYTAC目前面临的主要问题"
    + cite(7)
    + "，"
    "包括分子制备复杂、组织选择性不足以及临床前数据仍不够充分。",
    "Mamun等在综述中强调，"
    "eTPD药物需要同时解决分子设计、组织分布、免疫原性和制备可放大性等问题"
    + cite(16)
    + "。",
    "根据上述综述，截至相关文献发表时，"
    "LYTAC仍主要处于临床前研究阶段，尚未见正式发表的人体临床试验结果。",
    "因此，这一技术虽然值得持续关注，但距离真正成药还有较长距离。",
    "这一点与PROTAC等降解技术早期的发展路径有些相似"
    + cite(10, 17)
    + "，"
    "都需要先完成较扎实的机制研究和临床前验证工作。",
    "在安全性方面，文献讨论较多的问题包括脱靶降解和免疫原性。"
    "由于CI-M6PR等受体在多种组织都有表达，"
    "第一代LYTAC如果全身给药，理论上可能影响正常组织的蛋白稳态。",
    "含聚糖和蛋白的偶联物还可能被免疫系统识别，"
    "长期反复给药时的免疫反应需要在动物实验中认真评估。",
    "在药效学方面，LYTAC属于催化性较低、消耗性较强的模式。"
    "一个分子把靶蛋白和受体一起带入溶酶体后，"
    "往往自身也会被降解，因此体内有效浓度维持时间可能较短。",
    "从药物开发流程看，LYTAC目前仍缺少标准化的大分子偶联工艺和质量控制方法。",
    "相比之下，KineTAC、iLYTAC和EndoTag等基因编码平台在制备上更接近生物药生产路线"
    + cite(3, 12, 13)
    + "，"
    "但同样要面对表达、纯化和稳定性等问题。",
    "Caianiello等报道的MoDE-A策略说明，"
    "小分子也可能介导胞外蛋白降解"
    + cite(9)
    + "，"
    "这为未来开发更易制备的降解药物提供了另一种可能。",
    "Fang等的SignalTAC则表明，"
    "融合内吞信号肽亦可在不依赖外源聚糖配体的情况下实现膜蛋白降解"
    + cite(15)
    + "。",
    "近年来LYTAC领域最显著的变化在于技术平台的多样化。"
    "除了Banik等提出的聚糖-抗体偶联路线"
    + cite(1)
    + "，"
    "现已出现KineTAC、Apt-LYTAC、MoDE-A、iLYTAC、EndoTag、SignalTAC和病灶激活前体等多种形式"
    + cite(3, 4, 5, 6, 9, 12, 13, 15)
    + "。",
    "不同平台各有适用场景："
    "肝靶向疾病可能更适合GalNAc或Apt-LYTAC，"
    "肿瘤免疫可能关注PD-L1和HER2降解，"
    "神经退行性疾病则更强调血脑屏障穿透和局部激活。",
    "因此，评价一种LYTAC有没有前景，"
    "不能只看体外降解率，还要结合靶点、受体、给药途径和疾病模型来综合判断。",
    "展望未来，该领域的发展可能沿以下方向推进："
    "组织特异性受体的开发与验证、"
    "制备工艺简化与产物均一性提升，"
    "以及体内药效学、毒理学和药代动力学的系统研究。",
    "Lin等的综述和Mamun等的综述均指出，"
    "整合素、清道夫受体和细胞因子受体等亦值得作为组织选择性降解的受体类型"
    + cite(11, 16)
    + "。",
    "工艺与质量控制是LYTAC成药前必须解决的关键问题。"
    "抗体-聚糖偶联物往往存在偶联位点和偶联数不一致的问题，"
    "不同批次分子的降解活性可能出现波动。"
    "Pance等提出的KineTAC、Zhang等的iLYTAC以及Wu等和Li等采用的适体路线"
    + cite(3, 13, 6, 4)
    + "，"
    "在一定程度上是为了绕开复杂化学偶联带来的制备困难，"
    "但每种替代方案也会带来新的稳定性或免疫原性风险。",
    "该领域的研究论文通常同时涵盖化学结构、细胞实验与动物模型，"
    "对多学科知识背景具有较高要求。"
    "例如，理解Liu等的病灶激活LYTAC论文需掌握Aβ沉积病理、"
    "铜离子催化点击反应及血脑屏障递送等相关知识"
    + cite(5)
    + "；"
    "理解Kim等的酵母聚糖LYTAC则需了解M6P受体通路与糖缀合化学"
    + cite(14)
    + "。",
    "上述特点表明，新药研发是生物医学、药学与化学等多学科交叉融合的过程。",
    "LYTAC作为一项仍在快速演进的技术，"
    "其核心策略是借助溶酶体途径清除胞外及膜蛋白；"
    "其成功开发有赖于受体生物学、分子设计与疾病模型三个层面的协同优化。",
    "现有文献已提供较为丰富的概念验证数据，"
    "但距临床药物转化仍有明显差距。",
    "总体而言，LYTAC代表了靶向蛋白降解向胞外靶点拓展的重要尝试，"
    "其发展过程充分体现了基础生物医学研究对新药研发的推动作用。",
    "Bekes等在回顾PROTAC产业化历程时指出，"
    "靶向蛋白降解已从学术概念发展为具有明确临床路径的药物研发模式"
    + cite(10)
    + "。"
    "LYTAC目前虽尚未进入临床试验，"
    "但其在PD-L1、EGFR、Aβ等靶点上积累的概念验证数据，"
    "已为后续临床前开发奠定了基础"
    + cite(1, 4, 5)
    + "。",
    "未来研究可能集中在以下方面："
    "一是开发组织特异性更强的LTR配体或内吞诱导模块；"
    "二是建立可放大生产的制备与质量控制标准，"
    "尤其针对聚糖偶联物和基因编码降解蛋白；"
    "三是开展更系统的体内药效学、毒理学和药代动力学评价，"
    "为候选药物筛选提供依据"
    + cite(7, 8, 16)
    + "。",
    "Lin等和Chen等均强调，"
    "胞外蛋白降解技术的成功不仅取决于体外降解效率，"
    "还取决于靶点选择、受体表达谱、给药途径及疾病模型等多重因素"
    + cite(8, 11)
    + "。"
    "因此，LYTAC药物研发需要在分子化学、细胞生物学和疾病药理学之间保持紧密衔接，"
    "方能推动该技术从实验室研究走向临床转化。",
    "在分子均一性方面，抗体-聚糖偶联物常存在偶联位点和偶联度不一致的问题，"
    "导致不同批次产物活性波动。"
    "基因编码的KineTAC、iLYTAC和EndoTag在表达纯化后具有较好的序列均一性，"
    "但需关注蛋白药物的免疫原性和体内半衰期"
    + cite(3, 12, 13)
    + "。",
    "适体LYTAC虽合成简便，"
    "但体内可能面临核酸酶降解及肾脏快速清除等问题，"
    "需要通过化学修饰或载体递送加以改善"
    + cite(4, 6)
    + "。",
    "Fang等的SignalTAC以基因工程方式表达，"
    "通过融合内吞信号肽增强靶蛋白向溶酶体的分选，"
    "为膜蛋白降解提供了一种不依赖外源聚糖配体的替代方案"
    + cite(15)
    + "。",
    "Mamun等在综述结尾指出，"
    "eTPD领域当前面临的主要开放问题包括："
    "如何建立可重复的体外活性评价标准、"
    "如何预测不同组织中的脱靶降解风险、"
    "以及如何设计可放大生产的分子架构"
    + cite(16)
    + "。"
    "这些问题同样适用于LYTAC技术的后续开发。",
    "Zhao等在综述中强调，"
    "任何新兴降解平台均需经过充分的机制验证、"
    "药代毒理评估和制剂研究，方可进入临床开发"
    + cite(18)
    + "。"
    "对于LYTAC而言，"
    "溶酶体途径依赖性的确证、"
    "受体占用动力学表征及"
    "长期给药安全性评价将是临床前研究的重点内容。",
    "综合现有文献，"
    "LYTAC技术已在分子多样性、靶点覆盖面和疾病模型验证方面取得显著进展，"
    "但距离成为成熟治疗手段仍有较长路程。"
    "后续研究需在已有概念验证基础上，"
    "逐步解决制备工艺、组织选择性和临床证据不足等瓶颈问题，"
    "方能在靶向蛋白降解领域占据一席之地。",
    "在免疫相关疾病领域，"
    "胞外细胞因子、生长因子及免疫检查点分子均是潜在降解靶点。"
    "Chen等在综述中列举了细胞因子受体和生长因子受体等值得关注的对象"
    + cite(8)
    + "，"
    "但针对这类靶点的LYTAC公开数据仍相对有限。",
    "在制备工艺方面，"
    "化学合成多价M6P聚糖的高成本是制约LYTAC产业化的重要因素之一。"
    "Kim等通过酵母发酵制备人源兼容的M6P聚糖，"
    "并以点击化学完成与抗体的偶联，"
    "为降低聚糖配体制备门槛提供了可行方案"
    + cite(14)
    + "。",
    "Zhang等的iLYTAC则完全避开了聚糖化学合成，"
    "通过IGF2与CI-M6PR的天然配体-受体相互作用实现溶酶体靶向，"
    "且可在标准大肠杆菌表达系统中大量制备"
    + cite(13)
    + "。",
    "两种策略代表了LYTAC制备工艺优化的不同方向："
    "一是改进现有聚糖路线的可及性，"
    "二是开发不依赖聚糖化学的基因编码替代平台。",
    "从药效学角度看，"
    "Li等的共价适体LYTAC在肿瘤模型中诱导了比抗PD-L1抗体更强的T细胞浸润和细胞因子释放，"
    "提示蛋白降解策略可能在某些情境下优于单纯受体阻断"
    + cite(4)
    + "。",
    "Huang等的pLYTAC在PD-L1靶向降解后亦显示出增强的抗肿瘤活性，"
    "进一步支持了这一判断"
    + cite(12)
    + "。",
    "然而，这些积极结果均来自临床前模型，"
    "人体中的免疫微环境、受体表达水平及给药耐受性可能与动物模型存在显著差异，"
    "尚需更多研究加以验证。",
]

CONCLUSION = [
    "综上所述，LYTAC技术将靶向蛋白降解的应用范围从胞内蛋白拓展至胞外及膜蛋白，"
    "为传统上难以干预的靶点提供了新的药物研发策略。",
    "Banik等建立了LYTAC的基本技术框架"
    + cite(1)
    + "，"
    "Ahn等实现了肝组织特异性降解"
    + cite(2)
    + "，"
    "Pance等提供了基因编码的KineTAC蛋白降解平台"
    + cite(3)
    + "，"
    "Wu等、Li等、Zhang等和Huang等则分别从适体、共价适体、IGF2融合蛋白和EndoTag等方向拓展了技术路线"
    + cite(4, 6, 13, 12)
    + "。",
    "上述研究表明，成功的药物研发不仅依赖分子设计，"
    "还需深入理解受体分布、内吞动力学及溶酶体分选等生物学过程。",
    "在应用层面，LYTAC在PD-L1降解、肝靶向蛋白清除及Aβ清除等方面已获得动物实验支持，"
    "表明该技术具有一定可行性。",
    "然而，分子制备复杂、组织选择性不足及临床数据匮乏仍是制约其发展的主要瓶颈。",
    "LYTAC目前仍处于前沿探索阶段，"
    "尚不能等同于即将进入临床的成熟药物技术。",
    "随着机制研究的深入及新型平台的涌现，"
    "该领域有望在未来取得更多突破性进展。",
]

REFS = [
    "[1] Banik S M, Pedram K, Wisnovsky S, Riley N M, Bertozzi C R. Lysosome-targeting chimaeras for degradation of extracellular proteins[J]. Nature, 2020, 584(7820): 291-297. https://doi.org/10.1038/s41586-020-2545-9",
    "[2] Ahn G, Banik S M, Miller C L, Riley N M, Cochran J R, Bertozzi C R. LYTACs that engage the asialoglycoprotein receptor for targeted protein degradation[J]. Nature Chemical Biology, 2021, 17(9): 937-946. https://doi.org/10.1038/s41589-021-00770-1",
    "[3] Pance K, Gramespacher J A, Byrnes J R, et al. Modular cytokine receptor-targeting chimeras for targeted degradation of cell surface and extracellular proteins[J]. Nature Biotechnology, 2023, 41(2): 273-281. https://doi.org/10.1038/s41587-022-01456-2",
    "[4] Li Y, Liu X, Yu L, Huang X, Wang X, Han D, et al. Covalent LYTAC enabled by DNA aptamers for immune checkpoint degradation therapy[J]. Journal of the American Chemical Society, 2023, 145(45): 24506-24521. https://doi.org/10.1021/jacs.3c03899",
    "[5] Liu Z, Deng Q, Qin G, Yang J, Zhang H, Ren J, Qu X. Biomarker-activated multifunctional lysosome-targeting chimeras mediated selective degradation of extracellular amyloid fibrils[J]. Chem, 2023, 9(7): 2016-2038. https://doi.org/10.1016/j.chempr.2023.06.003",
    "[6] Wu Y, Lu Y, Li L, Deng K, Zhang S, Yang C, Zhu Z. Aptamer-LYTACs for targeted degradation of extracellular and membrane proteins[J]. Angewandte Chemie International Edition, 2023, 62(15): e202218106. https://doi.org/10.1002/anie.202218106",
    "[7] Li Y Y, Yang Y, Zhang R S, Ge R X, Xie S B. Targeted degradation of membrane and extracellular proteins with LYTACs[J]. Acta Pharmacologica Sinica, 2025, 46: 1-7. https://doi.org/10.1038/s41401-024-01364-y",
    "[8] Chen X, Zhou Y, Zhao Y, Tang W. Targeted degradation of extracellular secreted and membrane proteins[J]. Trends in Pharmacological Sciences, 2023, 44(11): 762-775. https://doi.org/10.1016/j.tips.2023.08.013",
    "[9] Caianiello D F, Zhang M, Ray J D, et al. Bifunctional small molecules that mediate the degradation of extracellular proteins[J]. Nature Chemical Biology, 2021, 17(9): 947-953. https://doi.org/10.1038/s41589-021-00851-1",
    "[10] Bekes M, Langley D R, Crews C M. PROTAC targeted protein degraders: the past is prologue[J]. Nature Reviews Drug Discovery, 2022, 21(3): 181-200. https://doi.org/10.1038/s41573-021-00371-6",
    "[11] Lin J, Jin J, Shen Y, Zhang L, Gong G, Bian H, Chen H, Nagle D G, Wu Y, Zhang W, Luan X. Emerging protein degradation strategies: expanding the scope to extracellular and membrane proteins[J]. Theranostics, 2021, 11(17): 8337-8349. https://doi.org/10.7150/thno.62686",
    "[12] Huang B, Abedi M, Coventry B, et al. Designed endocytosis-inducing proteins degrade targets and amplify signals[J]. Nature, 2024, 632(8024): 191-200. https://doi.org/10.1038/s41586-024-07948-2",
    "[13] Zhang B, Brahma R K, Zhu L, Feng J, Hu S, Qian L, et al. Insulin-like Growth Factor 2 (IGF2)-Fused Lysosomal Targeting Chimeras for Degradation of Extracellular and Membrane Proteins[J]. Journal of the American Chemical Society, 2023, 145(44): 24272-24283. https://doi.org/10.1021/jacs.3c08886",
    "[14] Kim S, Kang J Y, Bi A D, Seo J, Oh D B. Lysosome-Targeting Chimera Using Mannose-6-Phosphate Glycans Derived from Glyco-Engineered Yeast[J]. Bioconjugate Chemistry, 2025, 36(3): 424-436. https://doi.org/10.1021/acs.bioconjchem.4c00512",
    "[15] Fang T, Zheng Z, Li N, Zhang Y, Ma J, Yun C, Cai X. Lysosome-targeting chimeras containing an endocytic signaling motif trigger endocytosis and lysosomal degradation of cell-surface proteins[J]. Chemical Science, 2024, 15(42): 17652-17662. https://doi.org/10.1039/d4sc05093b",
    "[16] Mamun A A, Uzunparmak B, Crews C M. Targeted degradation of extracellular proteins: state of the art and diversity of degrader designs[J]. Journal of Hematology & Oncology, 2025, 18: 23. https://doi.org/10.1186/s13045-025-01703-4",
    "[17] Zhong G, Chang X, Xie W, Zhou X. Targeted protein degradation: advances in drug discovery and clinical practice[J]. Signal Transduction and Targeted Therapy, 2024, 9: 308. https://doi.org/10.1038/s41392-024-02004-x",
    "[18] Zhao L, Zhao J, Zhong K, Tong A, Jia D. Targeted protein degradation: mechanisms, strategies and application[J]. Signal Transduction and Targeted Therapy, 2022, 7: 113. https://doi.org/10.1038/s41392-022-00966-4",
]


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    add_heading_text(doc, "溶酶体靶向嵌合体（LYTAC）技术在药物研发中的研究进展", level=1)
    add_heading_text(doc, "——文献综述", level=1)

    add_heading_text(doc, "摘要", level=2)
    add_paragraph(doc, ABSTRACT, first_line_indent=False)
    kp = doc.add_paragraph()
    kp.paragraph_format.first_line_indent = Cm(0)
    set_run_font(kp.add_run("关键词：溶酶体靶向嵌合体，靶向蛋白降解，溶酶体靶向受体，药物研发，生物医学"), bold=True)

    add_heading_text(doc, "前言", level=2)
    for para in INTRO:
        add_paragraph(doc, para)

    add_heading_text(doc, "正文", level=2)

    add_heading_text(doc, "一、LYTAC的作用机制与分子设计基础", level=3)
    for para in BODY_SECTION1:
        add_paragraph(doc, para)
    add_figure_placeholder(
        doc,
        "图1  LYTAC介导的胞外及膜蛋白溶酶体降解机制示意图。"
        "（引自：Banik S M, et al. Lysosome-targeting chimaeras for degradation of extracellular proteins. Nature, 2020. [1]）",
    )

    add_heading_text(doc, "二、代表性研究进展", level=3)
    for para in BODY_SECTION2:
        add_paragraph(doc, para)
    add_figure_placeholder(
        doc,
        "图2  首批LYTAC分子的结构设计与靶蛋白降解验证（Banik et al., Nature, 2020 [1]）。"
        "图3  GalNAc-LYTAC的肝靶向降解（Ahn et al., Nat Chem Biol, 2021 [2]）。"
        "图4  KineTAC平台原理（Pance et al., Nat Biotechnol, 2023 [3]）。"
        "图5  酵母来源M6P聚糖LYTAC制备与PD-L1降解（Kim et al., Bioconjugate Chem, 2025 [14]）。",
    )

    add_heading_text(doc, "三、应用探索与面临的主要问题", level=3)
    for para in BODY_SECTION3:
        add_paragraph(doc, para)

    add_heading_text(doc, "总结与展望", level=2)
    for para in CONCLUSION:
        add_paragraph(doc, para)

    add_heading_text(doc, "参考文献", level=2)
    for ref in REFS:
        add_reference(doc, ref)

    output_path = "/workspace/LYTAC技术药物研发综述.docx"
    doc.save(output_path)

    body_text = "".join(BODY_SECTION1 + BODY_SECTION2 + BODY_SECTION3)
    body_count = count_chinese_chars(body_text)
    total_text = ABSTRACT + "".join(INTRO + BODY_SECTION1 + BODY_SECTION2 + BODY_SECTION3 + CONCLUSION)
    total_count = count_chinese_chars(total_text)
    print(f"Document saved to: {output_path}")
    print(f"正文汉字数（不含文献名标注）: {body_count}")
    print(f"全文汉字数（不含文献名标注）: {total_count}")


if __name__ == "__main__":
    main()
