#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""从综述全文 Markdown 生成 Word 文档。"""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

REFS = [
    "[1] Bekes M, Langley D R, Crews C M. PROTAC targeted protein degraders: the past is prologue[J]. Nature Reviews Drug Discovery, 2022, 21(3): 181-200.",
    "[2] Banik S M, Pedram K, Wisnovsky S, Riley N M, Bertozzi C R. Lysosome-targeting chimaeras for degradation of extracellular proteins[J]. Nature, 2020, 584(7820): 291-297.",
    "[3] Chen X, Zhou Y, Zhao Y, Tang W. Targeted degradation of extracellular secreted and membrane proteins[J]. Trends in Pharmacological Sciences, 2023, 44(11): 762-775.",
    "[4] Mamun A A, Uzunparmak B, Crews C M. Targeted degradation of extracellular proteins: state of the art and diversity of degrader designs[J]. Journal of Hematology & Oncology, 2025, 18: 23.",
    "[5] Lin J, Jin J, Shen Y, et al. Emerging protein degradation strategies: expanding the scope to extracellular and membrane proteins[J]. Theranostics, 2021, 11(17): 8337-8349.",
    "[6] Li Y Y, Yang Y, Zhang R S, Ge R X, Xie S B. Targeted degradation of membrane and extracellular proteins with LYTACs[J]. Acta Pharmacologica Sinica, 2025, 46: 1-7.",
    "[7] Huang B, Abedi M, Coventry B, et al. Designed endocytosis-inducing proteins degrade targets and amplify signals[J]. Nature, 2024, 632(8024): 191-200.",
    "[8] Zhao L, Zhao J, Zhong K, Tong A, Jia D. Targeted protein degradation: mechanisms, strategies and application[J]. Signal Transduction and Targeted Therapy, 2022, 7: 113.",
    "[9] Liu Z, Deng Q, Qin G, et al. Biomarker-activated multifunctional lysosome-targeting chimeras mediated selective degradation of extracellular amyloid fibrils[J]. Chem, 2023, 9(7): 2016-2038.",
    "[10] Ahn G, Banik S M, Miller C L, et al. LYTACs that engage the asialoglycoprotein receptor for targeted protein degradation[J]. Nature Chemical Biology, 2021, 17(9): 937-946.",
    "[11] Pance K, Gramespacher J A, Byrnes J R, et al. Modular cytokine receptor-targeting chimeras for targeted degradation of cell surface and extracellular proteins[J]. Nature Biotechnology, 2023, 41(2): 273-281.",
    "[12] Li Y, Liu X, Yu L, et al. Covalent LYTAC enabled by DNA aptamers for immune checkpoint degradation therapy[J]. Journal of the American Chemical Society, 2023, 145(45): 24506-24521.",
    "[13] Wu Y, Lu Y, Li L, et al. Aptamer-LYTACs for targeted degradation of extracellular and membrane proteins[J]. Angewandte Chemie International Edition, 2023, 62(15): e202218106.",
    "[14] Caianiello D F, Zhang M, Ray J D, et al. Bifunctional small molecules that mediate the degradation of extracellular proteins[J]. Nature Chemical Biology, 2021, 17(9): 947-953.",
    "[15] Zhang B, Brahma R K, Zhu L, et al. Insulin-like Growth Factor 2 (IGF2)-Fused Lysosomal Targeting Chimeras for Degradation of Extracellular and Membrane Proteins[J]. Journal of the American Chemical Society, 2023, 145(44): 24272-24283.",
    "[16] Kim S, Kang J Y, Bi A D, et al. Lysosome-Targeting Chimera Using Mannose-6-Phosphate Glycans Derived from Glyco-Engineered Yeast[J]. Bioconjugate Chemistry, 2025, 36(3): 424-436.",
    "[17] Fang T, Zheng Z, Li N, et al. Lysosome-targeting chimeras containing an endocytic signaling motif trigger endocytosis and lysosomal degradation of cell-surface proteins[J]. Chemical Science, 2024, 15(42): 17652-17662.",
    "[18] Zhong G, Chang X, Xie W, Zhou X. Targeted protein degradation: advances in drug discovery and clinical practice[J]. Signal Transduction and Targeted Therapy, 2024, 9: 308.",
]


def set_run_font(run, chinese_font="仿宋", western_font="Times New Roman", size_pt=12, bold=False):
    run.font.name = western_font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), chinese_font)


def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    set_run_font(p.add_run(text), size_pt=sizes.get(level, 12), bold=True)


def add_body_paragraph(doc, text, first_line_indent=True, center=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0.74) if first_line_indent and not center else Cm(0)
    if text.startswith("【本节小结】"):
        set_run_font(p.add_run("【本节小结】"), bold=True)
        set_run_font(p.add_run(text[len("【本节小结】"):]))
    elif text.startswith("【此处插入"):
        set_run_font(p.add_run(text), size_pt=12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif text.startswith("图") and "引自" in text:
        set_run_font(p.add_run(text), size_pt=10.5)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        set_run_font(p.add_run(text))
    return p


def add_reference(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0.74)
    pf.hanging_indent = Cm(0.74)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(3)
    set_run_font(p.add_run(text), size_pt=10.5)


def parse_markdown(md_text):
    """解析 markdown，返回 (block_type, text) 列表。"""
    blocks = []
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line or line == "---":
            i += 1
            continue
        if line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
            i += 1
            continue
        if line.startswith("## "):
            blocks.append(("h3", line[3:].strip()))
            i += 1
            continue
        if line.startswith("【") or (line.startswith("图") and "引自" in line):
            blocks.append(("special", line))
            i += 1
            continue
        para_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt == "---" or nxt.startswith("#") or nxt.startswith("##") or nxt.startswith("【") or (nxt.startswith("图") and "引自" in nxt):
                break
            para_lines.append(nxt)
            i += 1
        blocks.append(("p", "".join(para_lines)))
    return blocks


def build_document(md_path, output_path):
    md_text = Path(md_path).read_text(encoding="utf-8")
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    blocks = parse_markdown(md_text)
    in_refs = False
    summary_next = False
    for kind, text in blocks:
        if text == "参考文献":
            in_refs = True
            add_heading(doc, text, level=2)
            continue
        if in_refs:
            if text.startswith("["):
                add_reference(doc, text)
            continue
        if kind == "h1":
            if text in ("摘要", "前言", "正文", "总结与展望", "参考文献"):
                add_heading(doc, text, level=2)
            elif text in ("溶酶体靶向嵌合体（LYTAC）技术在药物研发中的研究进展", "——文献综述"):
                add_heading(doc, text, level=1)
            else:
                add_heading(doc, text, level=3)
        elif kind == "h3":
            if text == "本节小结":
                summary_next = True
            else:
                add_heading(doc, text, level=3)
        elif kind == "special":
            if text.startswith("【此处插入"):
                add_body_paragraph(doc, "【此处插入示意图】", center=True, first_line_indent=False)
                cap = re.sub(r"^【此处插入图\d+】", "", text).strip()
                if cap:
                    add_body_paragraph(doc, cap, center=True, first_line_indent=False)
            else:
                add_body_paragraph(doc, text, center=True, first_line_indent=False)
        else:
            if text.startswith("关键词："):
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0)
                set_run_font(p.add_run(text), bold=True)
            elif summary_next:
                add_body_paragraph(doc, "【本节小结】" + text, first_line_indent=True)
                summary_next = False
            else:
                add_body_paragraph(doc, text)

    doc.save(output_path)
    print(f"已生成: {output_path}")


if __name__ == "__main__":
    build_document("/workspace/综述全文.md", "/workspace/LYTAC技术药物研发综述.docx")
