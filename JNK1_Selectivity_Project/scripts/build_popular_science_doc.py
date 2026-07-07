#!/usr/bin/env python3
"""Build popular-science Word document and data_tables folder from project report."""

from __future__ import annotations

import json
import shutil
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "popular_science"
DATA_DIR = OUT / "data_tables"
FIG_DIR = OUT / "figures"
DOC_PATH = OUT / "JNK1筛选项目通俗解读报告.docx"

plt.rcParams["font.sans-serif"] = ["DejaVu Sans", "Arial Unicode MS", "SimHei", "sans-serif"]
plt.rcParams["axes.unicode_minus"] = False

# --- Data tables to copy / synthesize -----------------------------------------

COPY_MAP = {
    "01_文献benchmark化合物.csv": "data/benchmarks/literature_benchmarks.csv",
    "02_ML模型性能comparison.json": "results/model_comparison/comparison.json",
    "03_ML外部decoy验证指标.json": "results/ml_external_validation/ml_external_validation_metrics_9bd8.json",
    "04_再对接验证redocking.csv": "results/docking_validation/redocking_summary_7725.csv",
    "05_benchmark对接差值.csv": "results/docking_validation/benchmark_deltas_51c1.csv",
    "06_benchmark方向混淆矩阵.csv": "results/docking_validation/direction_confusion_27c3.csv",
    "07_benchmark_MMGBSA标定.csv": "results/docking_validation/benchmark_mmgbsa_calibration.csv",
    "08_Gly87自检.csv": "results/docking_validation/gly87_selfcheck_16be.csv",
    "09_各亚型排序相关.csv": "results/docking_validation/isoform_rank_correlations_299a.csv",
    "10_选择性标签统计.csv": "results/similarity/sel_class_counts.csv",
    "11_采购清单purchase_after_md.csv": "data/purchase/purchase_after_md.csv",
    "12_2231延伸MD_RMSD分位数.csv": "results/md_2231_200ns/tables/09_production_rmsd_percentiles.csv",
    "13_2231延伸MD_MMGBSA分量.csv": "results/md_2231_200ns/tables/14_mm_gbsa_components.csv",
    "14_ML阈值校准threshold.json": "results/calibration/threshold_recommendation.json",
    "15_ML虚拟筛选demo_screening.json": "results/screening_v2/screening_report.json",
}


def synthesize_tables() -> None:
  """Create summary tables documented in JNK1_PROJECT_REPORT.md."""
  pd.DataFrame([
    {"阶段": "ML初筛后对接库(F0)", "数量": 4983, "数据来源": "md_shortlist_report_23c8.md"},
    {"阶段": "Glide XP VSW有效记录", "数量": 4979, "数据来源": "JNK1_SELECTIVITY_FINAL_REPORT_41d9.md"},
    {"阶段": "MD短名单 F1∧F2(ADMET前)", "数量": 157, "数据来源": "md_shortlist_report_23c8.md"},
    {"阶段": "MD shortlist(ADMET后)", "数量": 25, "数据来源": "同上"},
    {"阶段": "MD pose QC输入", "数量": 16, "数据来源": "MD_QC_report_cf26.md"},
    {"阶段": "最终采购推荐", "数量": 10, "数据来源": "data/purchase/purchase_after_md.csv"},
  ]).to_csv(DATA_DIR / "00_端到端漏斗汇总.csv", index=False, encoding="utf-8-sig")

  pd.DataFrame([
    {"阶段": "ML F1后", "数量": 4983, "条件": "p_family ≥ 6.0"},
    {"阶段": "VSW有效", "数量": 4979, "条件": "三isoform均有XP分"},
    {"阶段": "pass_pose", "数量": 3234, "条件": "Glide pose质量门"},
    {"阶段": "pass_potency", "数量": 1681, "条件": "score_JNK1 ≤ -7.43"},
    {"阶段": "pass_selectivity(遗留)", "数量": 233, "条件": "Δsel_dock>0 且 Δsel_MMGBSA≥2"},
    {"阶段": "has_selectivity_contact", "数量": 63, "条件": "铰链H-bond代理+Δsel启发式"},
  ]).to_csv(DATA_DIR / "16_VSW选择性探索各阶段数量.csv", index=False, encoding="utf-8-sig")

  pd.DataFrame([
    {"Tier": "Tier 1′", "数量": 57, "条件": "pose+potency+pass_selectivity+contact"},
    {"Tier": "Tier 2", "数量": 92, "条件": "pose+potency+pass_selectivity"},
    {"Tier": "Tier 3", "数量": 1191, "条件": "pose+potency"},
    {"Tier": "Tier 0", "数量": 3639, "条件": "未达Tier3"},
  ]).to_csv(DATA_DIR / "17_Tier分布.csv", index=False, encoding="utf-8-sig")

  pd.DataFrame([
    {"阶段": "输入(F0后)", "数量": 4983},
    {"阶段": "F1 pose QC通过", "数量": 3125},
    {"阶段": "F2活性+配体效率通过", "数量": 182},
    {"阶段": "F1∧F2通过", "数量": 157},
    {"阶段": "F7 ADMET剔除", "数量": 9},
    {"阶段": "ADMET后shortlist", "数量": 25},
    {"阶段": "进入MD", "数量": 16},
    {"阶段": "采购推荐", "数量": 10},
  ]).to_csv(DATA_DIR / "18_MD短名单漏斗.csv", index=False, encoding="utf-8-sig")

  pd.DataFrame([
    {"序号": 0, "阶段": "ML F1后对接库(F0)", "数量": 4983, "评判指标与门槛": "p_family ≥ 6.0", "指标类别": "ML活性召回", "用于采购链": "是", "数据来源": "md_shortlist_report_23c8.md", "备注": ""},
    {"序号": 1, "阶段": "VSW有效", "数量": 4979, "评判指标与门槛": "三isoform均有Glide XP gscore", "指标类别": "对接有效性", "用于采购链": "是", "数据来源": "JNK1_SELECTIVITY_FINAL_REPORT_41d9.md", "备注": "4个对接失败"},
    {"序号": 2, "阶段": "pass_pose(VSW pose门)", "数量": 3234, "评判指标与门槛": "Glide pose质量布尔门", "指标类别": "pose可信度(对接层A)", "用于采购链": "间接", "数据来源": "同上", "备注": "与MD-F1部分重叠"},
    {"序号": 3, "阶段": "pass_potency(单活性门)", "数量": 1681, "评判指标与门槛": "score_JNK1 ≤ -7.43", "指标类别": "JNK1活性单点", "用于采购链": "是(分量)", "数据来源": "同上", "备注": "可与pass_pose独立计数"},
    {"序号": 4, "阶段": "MD-F1 pose QC", "数量": 3125, "评判指标与门槛": "更严Glide pose质量门", "指标类别": "pose可信度(对接层B)", "用于采购链": "是", "数据来源": "md_shortlist_report_23c8.md", "备注": "比pass_pose少109"},
    {"序号": 5, "阶段": "MD-F2活性双门槛", "数量": 182, "评判指标与门槛": "score_JNK1 ≤ -7.43 且 MMGBSA_JNK1 ≤ -51.6", "指标类别": "JNK1活性双点", "用于采购链": "是", "数据来源": "同上", "备注": "非Δsel"},
    {"序号": "5a", "阶段": "F2∩pass_pose(推断)", "数量": 165, "评判指标与门槛": "双门槛∩pass_pose", "指标类别": "pose+活性", "用于采购链": "推断", "数据来源": "未入库", "备注": "不得编造ID"},
    {"序号": 6, "阶段": "F1∧F2(MD短名单ADMET前)", "数量": 157, "评判指标与门槛": "MD-F1 且 MD-F2", "指标类别": "pose+活性", "用于采购链": "是", "数据来源": "md_shortlist_report_23c8.md", "备注": ""},
    {"序号": "—", "阶段": "pass_selectivity(探索)", "数量": 233, "评判指标与门槛": "Δsel_dock>0 且 Δsel_MMGBSA≥2", "指标类别": "选择性探索", "用于采购链": "否", "数据来源": "JNK1_SELECTIVITY_FINAL_REPORT_41d9.md", "备注": "benchmark否定"},
  ]).to_csv(DATA_DIR / "26_对接后筛选漏斗.csv", index=False, encoding="utf-8-sig")

  pd.DataFrame([
    {"指标": "Spearman(Δsel_dock_vsw, -ΔpIC50_sel)", "数值": 0.750, "n": 7},
    {"指标": "Spearman(Δsel_mmgbsa_vsw, -ΔpIC50_sel)", "数值": 0.786, "n": 7},
    {"指标": "方向准确率(docking,VSW PDB)", "数值": "43%(3/7)", "n": 7},
    {"指标": "方向准确率(MM-GBSA,VSW PDB)", "数值": "43%(3/7)", "n": 7},
    {"指标": "方向准确率(docking,ensemble归档)", "数值": "29%(2/7)", "n": 7},
  ]).to_csv(DATA_DIR / "19_benchmark定量结果.csv", index=False, encoding="utf-8-sig")

  pd.DataFrame([
    {"ID": 2231, "组": "G2", "JNK1": 0.91, "JNK2": 0.00, "JNK3": 0.10},
    {"ID": 2157, "组": "G1", "JNK1": 0.85, "JNK2": 0.46, "JNK3": 0.02},
    {"ID": 2232, "组": "G1", "JNK1": 1.00, "JNK2": 1.00, "JNK3": 0.04},
    {"ID": 690, "组": "G1", "JNK1": 1.00, "JNK2": 0.51, "JNK3": 0.77},
    {"ID": 1280, "组": "G2", "JNK1": 0.00, "JNK2": 0.58, "JNK3": 0.43},
    {"ID": 4795, "组": "G2", "JNK1": 0.04, "JNK2": 0.04, "JNK3": 0.49},
  ]).to_csv(DATA_DIR / "20_采购分子铰链占有率.csv", index=False, encoding="utf-8-sig")

  pd.DataFrame([
    {"ID": 2231, "JNK1": 0.48, "JNK2": 1.17, "JNK3": 0.66},
    {"ID": 2157, "JNK1": 0.49, "JNK2": 1.13, "JNK3": 0.35},
    {"ID": 690, "JNK1": 0.72, "JNK2": 1.98, "JNK3": 0.61},
    {"ID": 2232, "JNK1": 0.57, "JNK2": 0.29, "JNK3": 1.31},
  ]).to_csv(DATA_DIR / "21_采购分子配体RMSD中位数_Angstrom.csv", index=False, encoding="utf-8-sig")

  pd.DataFrame([
    {"亚型": "JNK1", "化合物数": 444, "Holdout_R2": 0.697, "Holdout_Spearman": 0.858, "Holdout_n": 31},
    {"亚型": "JNK2", "化合物数": 610, "Holdout_R2": 0.574, "Holdout_Spearman": 0.780, "Holdout_n": 67},
    {"亚型": "JNK3", "化合物数": 1147, "Holdout_R2": 0.774, "Holdout_Spearman": 0.869, "Holdout_n": 98},
  ]).to_csv(DATA_DIR / "22_ML模型性能汇总.csv", index=False, encoding="utf-8-sig")

  pd.DataFrame([
    {"组分": "Decoys(Taosu)", "n": 10000, "标签": "假定无活性"},
    {"组分": "Benchmarks", "n": 9, "标签": "已知活性"},
    {"组分": "ChEMBL actives", "n": 1210, "标签": "已知活性"},
  ]).to_csv(DATA_DIR / "23_外部decoy验证设计.csv", index=False, encoding="utf-8-sig")

  pd.DataFrame([
    {"指标": "Sensitivity(recall)", "数值": "99.3%"},
    {"指标": "Specificity", "数值": "4.7%"},
    {"指标": "Decoy FPR", "数值": "95.3%"},
    {"指标": "Precision", "数值": "11.3%"},
    {"指标": "ROC-AUC", "数值": 0.876},
    {"指标": "EF1%", "数值": 9.20},
  ]).to_csv(DATA_DIR / "24_外部decoy验证指标汇总.csv", index=False, encoding="utf-8-sig")


def setup_data_folder() -> None:
  DATA_DIR.mkdir(parents=True, exist_ok=True)
  FIG_DIR.mkdir(parents=True, exist_ok=True)
  for dst, src in COPY_MAP.items():
    shutil.copy2(ROOT / src, DATA_DIR / dst)
  synthesize_tables()
  readme = """# 数据表格文件夹

本文件夹收录 JNK1/2/3 计算筛选项目报告中引用的全部可复现数据表。
所有数值均可回溯至仓库内原始文件；合成汇总表（00、16–24 号）的字段与
`docs/JNK1_PROJECT_REPORT.md`（v2.8）正文一致。

## 文件索引

| 文件名 | 说明 | 原始路径 |
|--------|------|----------|
"""
  for dst, src in sorted(COPY_MAP.items()):
    readme += f"| {dst} | 见文件名 | `{src}` |\n"
  readme += """
| 00_端到端漏斗汇总.csv | 摘要漏斗各阶段数量 | 报告§摘要 |
| 16–24 号 CSV | 报告正文汇总表 | 报告对应章节 |

生成脚本：`scripts/build_popular_science_doc.py`
"""
  (DATA_DIR / "README.md").write_text(readme, encoding="utf-8")


def make_figures() -> dict[str, Path]:
  paths = {}

  funnel = pd.read_csv(DATA_DIR / "00_端到端漏斗汇总.csv")
  fig, ax = plt.subplots(figsize=(10, 5))
  y = funnel["数量"][::-1]
  labels = funnel["阶段"][::-1]
  ax.barh(range(len(y)), y, color="#4C72B0")
  ax.set_yticks(range(len(labels)))
  ax.set_yticklabels(labels, fontsize=9)
  ax.set_xlabel("化合物数量")
  ax.set_title("端到端筛选漏斗（有数据支撑的阶段）")
  for i, v in enumerate(y):
    ax.text(v + 30, i, str(int(v)), va="center", fontsize=9)
  fig.tight_layout()
  p = FIG_DIR / "fig01_funnel.png"
  fig.savefig(p, dpi=150)
  plt.close(fig)
  paths["funnel"] = p

  redock = pd.read_csv(DATA_DIR / "04_再对接验证redocking.csv")
  fig, ax = plt.subplots(figsize=(7, 4))
  ax.bar(redock["PDB_ID"], redock["rmsd_A"], color=["#55A868" if x else "#C44E52" for x in redock["pass(<2A)"]])
  ax.axhline(2.0, color="red", linestyle="--", label="阈值 2.0 Å")
  ax.set_ylabel("RMSD (Å)")
  ax.set_title("共晶再对接验证（5/5 通过）")
  ax.legend()
  fig.tight_layout()
  p = FIG_DIR / "fig02_redocking.png"
  fig.savefig(p, dpi=150)
  plt.close(fig)
  paths["redock"] = p

  bench = pd.read_csv(DATA_DIR / "06_benchmark方向混淆矩阵.csv")
  fig, ax = plt.subplots(figsize=(6, 4))
  colors = ["#55A868" if m else "#C44E52" for m in bench["direction_match"]]
  ax.bar(bench["name"], bench["delta_sel_dock"], color=colors)
  ax.axhline(0, color="black", linewidth=0.8)
  ax.set_ylabel("Δsel_dock (kcal/mol 量级)")
  ax.set_title("文献对照：对接选择性方向预测（绿=正确，红=错误）")
  plt.xticks(rotation=30, ha="right")
  fig.tight_layout()
  p = FIG_DIR / "fig03_benchmark_direction.png"
  fig.savefig(p, dpi=150)
  plt.close(fig)
  paths["bench_dir"] = p

  rmsd2231 = pd.read_csv(DATA_DIR / "12_2231延伸MD_RMSD分位数.csv")
  lig = rmsd2231[rmsd2231["Metric"] == "Ligand"]
  fig, ax = plt.subplots(figsize=(6, 4))
  systems = lig["System"].tolist()
  med = lig["median"].tolist()
  desmond = [0.48, 1.17, 0.66]
  x = np.arange(len(systems))
  w = 0.35
  ax.bar(x - w/2, med, w, label="Amber 200ns 中位数")
  ax.bar(x + w/2, desmond, w, label="Desmond 短MD 中位数(报告§6.4.2)")
  ax.set_xticks(x)
  ax.set_xticklabels(systems)
  ax.set_ylabel("配体 RMSD (Å)")
  ax.set_title("化合物 2231：三 isoform 配体稳定性对比")
  ax.legend(fontsize=8)
  fig.tight_layout()
  p = FIG_DIR / "fig04_2231_rmsd.png"
  fig.savefig(p, dpi=150)
  plt.close(fig)
  paths["2231_rmsd"] = p

  hinge = pd.read_csv(DATA_DIR / "20_采购分子铰链占有率.csv")
  fig, ax = plt.subplots(figsize=(8, 4))
  ids = hinge["ID"].astype(str)
  x = np.arange(len(ids))
  w = 0.25
  ax.bar(x - w, hinge["JNK1"], w, label="JNK1")
  ax.bar(x, hinge["JNK2"], w, label="JNK2")
  ax.bar(x + w, hinge["JNK3"], w, label="JNK3")
  ax.axhline(0.3, color="gray", linestyle="--", label="pass 阈值 0.3")
  ax.set_xticks(x)
  ax.set_xticklabels(ids)
  ax.set_ylabel("铰链氢键占有率 (0–1)")
  ax.set_title("采购分子：三亚型铰链氢键占有率")
  ax.legend(fontsize=8)
  fig.tight_layout()
  p = FIG_DIR / "fig05_hinge_occupancy.png"
  fig.savefig(p, dpi=150)
  plt.close(fig)
  paths["hinge"] = p

  metrics = pd.read_csv(DATA_DIR / "19_benchmark定量结果.csv")
  fig, ax = plt.subplots(figsize=(7, 4))
  dirs = ["VSW对接", "VSW MM-GBSA", "ensemble对接"]
  acc = [43, 43, 29]
  ax.bar(dirs, acc, color="#4C72B0")
  ax.axhline(55, color="red", linestyle="--", label="期望阈值 55%")
  ax.set_ylabel("方向准确率 (%)")
  ax.set_title("计算选择性方向预测准确率（均不达标）")
  ax.set_ylim(0, 70)
  ax.legend()
  fig.tight_layout()
  p = FIG_DIR / "fig06_direction_accuracy.png"
  fig.savefig(p, dpi=150)
  plt.close(fig)
  paths["dir_acc"] = p

  return paths


def set_cell_font(cell, size=9):
  for p in cell.paragraphs:
    for r in p.runs:
      r.font.size = Pt(size)


def add_table(doc: Document, df: pd.DataFrame, title: str, max_rows: int | None = None):
  doc.add_heading(title, level=3)
  if max_rows:
    df = df.head(max_rows)
  table = doc.add_table(rows=1, cols=len(df.columns))
  table.style = "Table Grid"
  hdr = table.rows[0].cells
  for i, c in enumerate(df.columns):
    hdr[i].text = str(c)
    set_cell_font(hdr[i], 9)
  for _, row in df.iterrows():
    cells = table.add_row().cells
    for i, v in enumerate(row):
      cells[i].text = str(v) if pd.notna(v) else ""
      set_cell_font(cells[i], 8)
  doc.add_paragraph()


def add_image(doc: Document, path: Path, caption: str, width_cm=15):
  doc.add_picture(str(path), width=Cm(width_cm))
  p = doc.add_paragraph(caption)
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  if p.runs:
    p.runs[0].italic = True
  doc.add_paragraph()


def build_word(fig_paths: dict[str, Path]) -> None:
  doc = Document()
  style = doc.styles["Normal"]
  style.font.name = "宋体"
  style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
  style.font.size = Pt(11)

  title = doc.add_heading("JNK1/2/3 抑制剂计算机辅助筛选项目\n通俗解读报告", 0)
  title.alignment = WD_ALIGN_PARAGRAPH.CENTER
  doc.add_paragraph("版本：2.7（面向非 CADD 专业读者）")
  doc.add_paragraph("数据来源：项目仓库 JNK1_Selectivity_Project，所有数值可回溯至 data_tables/ 文件夹")
  doc.add_paragraph("技术报告原文：docs/JNK1_PROJECT_REPORT.md")
  doc.add_page_break()

  sections = [
    ("写给完全不懂计算机辅助药物设计（CADD）的读者", """
本报告用通俗语言说明：我们如何从约五千个候选分子出发，经过机器学习、分子对接、成药性评估和分子动力学模拟，最终选出 10 个分子去做湿实验（试管酶学测定）。

需要事先明确的三件事：
第一，这是一套「计算预筛选」流程，目的是在花钱合成/购买化合物之前，用计算机尽量排除明显不靠谱的分子。
第二，本项目最初希望找到「主要抑制 JNK1、少抑制 JNK2/JNK3」的选择性分子，但计算验证表明：现有方法无法可靠判断亚型方向，因此策略已调整为寻找「可能与 JNK 家族结合的候选分子」，选择性只能交给实验。
第三，文中每一个数字都有原始数据文件支撑，汇总表见本报告配套文件夹 data_tables/。
"""),
    ("核心术语解释（阅读前请先浏览）", """
• JNK（c-Jun N-terminal kinase）：一类与炎症、纤维化等相关的蛋白激酶，有 JNK1、JNK2、JNK3 三个亚型（同工酶）。
• 亚型 / isoform：同一基因家族下结构相似但功能略有差别的蛋白版本。
• 抑制剂 / 配体 / 小分子：有望结合蛋白并调节其活性的化合物。
• IC50：使酶活性下降一半所需化合物浓度，越小通常活性越强。
• pIC50 = −log10(IC50/M)：活性常用对数刻度，数值越大活性越强。
• ChEMBL：公开的药物化学数据库，收录化合物生物活性（参考文献 [1]）。
• 虚拟筛选：用计算机从大量分子中快速缩小候选范围，而非逐个做实验。
• 机器学习（ML）：让计算机从已知「分子结构—活性」数据中学习规律，再预测新分子活性。
• XGBoost：一种常用的机器学习算法（参考文献 [9]）。
• 分子对接（docking）：把小分子「摆放」到蛋白结合口袋，用打分函数估计结合强弱（Glide，参考文献 [7]）。
• PDB：蛋白三维结构数据库编号；本项目用 3ELJ(JNK1)、3E7O(JNK2)、3TTI(JNK3) 等共晶结构。
• RMSD：两个分子结构相差多少埃（Å），越小越相似；再对接 RMSD < 2 Å 表示计算能复现实验结构。
• MM-GBSA：在对接 pose 基础上估算结合自由能（本项目用 Schrödinger Prime 实现）。
• ADMET：吸收、分布、代谢、排泄、毒性等成药性指标；本项目用 QikProp 预测。
• 分子动力学（MD）：模拟分子在溶液中随时间运动，观察结合是否稳定。
• Decoy（诱饵分子）：假定无活性的对照分子，用于检验筛选是否会把「噪音」当成活性。
• pan-JNK：对三个亚型都有一定抑制活性的「广谱」分子，而非只偏向 JNK1。
"""),
    ("一、项目背景：我们在找什么药？", """
JNK 三个亚型由不同基因编码（JNK1/MAPK8、JNK2/MAPK9、JNK3/MAPK10）。JNK1 与特发性肺纤维化（IPF）、非酒精性脂肪性肝炎（NASH）等疾病相关。

临床上已有 JNK 相关药物研究，例如 CC-90001（参考文献 [2]）在细胞层面偏向 JNK1 功能抑制；SP600125（参考文献 [3]）是经典的 pan-JNK 研究工具药。本项目还纳入文献报道的强 JNK1 抑制剂 E1（参考文献 [4]）作为对照。

项目初期目标：通过计算筛选找到 preferentially 抑制 JNK1 的候选分子。
项目后期结论：计算无法可靠区分亚型方向，因此把候选定位为「JNK 家族结合剂」，用同批次 JNK1/2/3 酶学 IC50 实验做最终判断（参考文献 [8] 综述了 JNK 成药挑战）。
"""),
    ("二、整体筛选流程（单一主线漏斗）", """
从 Enamine 商业库约 5000 个分子出发，经过以下主链路：

步骤 1 — 机器学习粗筛：用三个亚型活性模型判断分子是否「可能有 JNK 家族活性」，保留 4983 个。
步骤 2 — Glide XP 对接：把每个分子分别对接到 JNK1/JNK2/JNK3 三个蛋白结构，得到 4979 个有效结果。
步骤 3 — MD 短名单筛选：按 pose 质量、JNK1 活性门槛、成药性、ADMET 筛选 → 157 → 25 → 16 个做 MD → 最终 10 个采购。
步骤 4 — 湿实验：同批次 JNK1/JNK2/JNK3 酶学 IC50。

主线漏斗不使用 Δsel、pass_selectivity 或 Tier 分级。项目早期曾尝试用对接能量差 Δsel 判断 JNK1 选择性，但文献 benchmark 证明该指标方向准确率仅 43%（VSW 单结构）或 29%（双结构平均），不能用于采购决策。这些失败的探索尝试完整保留于第五章，供回顾分析。
"""),
    ("三、机器学习：如何粗筛「可能有活性」的分子？", """
3.1 训练数据与模型
从 ChEMBL（参考文献 [1]）收集 JNK1/2/3 活性数据，清洗后：JNK1 444 个、JNK2 610 个、JNK3 1147 个化合物。
采用 XGBoost（参考文献 [9]）分别训练三个回归模型，预测 pActivity（活性对数刻度）。
数据划分采用 Murcko 骨架分组：化学结构相似的分子不会同时出现在训练集和测试集，避免「背答案」。

独立 holdout 测试集性能（数据表 22_ML模型性能汇总.csv）：
JNK1 R²=0.697，JNK2 R²=0.574，JNK3 R²=0.774。
说明：模型对「这个分子强不强」有一定预测力，但对「偏向哪个亚型」不可靠。

3.2 为什么用 F1 门槛 p_family ≥ 6.0？
九个文献对照分子在 6.0 阈值下全部通过（9/9），目的是高召回——不让已知活性分子被漏掉。
外部 decoy 验证（10,000 个 Taosu 随机分子 + 1,219 个已知活性）显示：decoy 假阳性率 95.3%，特异性仅 4.7%，ROC-AUC=0.876，EF1%=9.20（数据表 23–24）。
解读：F1 是「宽网」，负责不漏真活性；真正去噪音靠综合排序分 final_score（含 QED、合成可及性 SA 等）和后续对接。

3.3 为什么 ML 不能判断亚型选择性？
ChEMBL 中明确标注 JNK1-selective 的分子仅 8 个（数据表 10），数据太少无法训练可靠分类器（测试 F1=0）。
对 E1、TCS JNK 6O 等对照，ML 预测的最高活性亚型与实验不符（数据表 05）。
"""),
    ("四、分子对接：如何把分子「放进」蛋白口袋？", """
4.1 软件与结构
使用 Schrödinger Glide XP 模式（参考文献 [7]）。每个亚型选一个主 PDB：JNK1→3ELJ，JNK2→3E7O，JNK3→3TTI。
打分采用 XP 终分 gscore，越负表示预测结合越强。

4.2 为什么要做共晶再对接？
在筛选大量分子之前，先拿共晶配体重新对接回自己的蛋白结构。若 RMSD < 2 Å，说明蛋白结构处理和对接参数可信。
结果 5/5 全部通过（数据表 04、图 2）。注意：这只能证明「结构准备正确」，不能证明「能预测选择性」。

4.3 选择性指标 Δsel 的含义与失败
Δsel_dock = min(score_JNK2, score_JNK3) − score_JNK1。若 > 0，计算上「更偏向 JNK1」。
对 9 个有实验 IC50 的文献分子做回顾：Spearman 秩相关尚可，但离散方向准确率仅 43%（3/7），低于 55% 可用线（数据表 19、图 6）。
因此：233 个 pass_selectivity 和 57 个 Tier 1′ 属于选择性探索统计（详见第五章），不用于 MD 进门。

4.4 对接后主线筛选漏斗（数据表 26，技术报告 §3.5）
4979 个有效对接后，采购链逐层为：
• pass_pose 3234：VSW Glide pose 质量门（探索统计，较宽）
• pass_potency 1681：score_JNK1 ≤ −7.43（单活性门，可与 pose 独立计数）
• MD-F1 3125：更严 pose QC（比 pass_pose 少 109 个）→ MD-F2 182：Glide + MM-GBSA 双活性门槛 → F1∧F2 共 157 个进入 ADMET
活性门槛 −7.43 来自 8 个非共价 benchmark 在 3ELJ 的 Glide 中位数；MMGBSA_JNK1 ≤ −51.6 为 JNK1 单点能量门，与 Δsel 无关。

4.5 两层 pose QC 的区别
对接阶段：pass_pose（3234，较宽）与 MD-F1（3125，较严）均为 Glide pose 质量评判，精确 Schrödinger 规则待归档。
MD 阶段：另用配体 RMSD ≤ 3 Å 与铰链氢键占有率 ≥ 30%（§7），与对接层数值不可直接对比。
"""),
    ("五、曾尝试的选择性策略与失败原因", """
5.0 VSW 对接后选择性探索（未用于采购）
项目在对接完成后曾尝试多种选择性标签，均经 benchmark 标定后从决策链移除，但统计数字保留供回顾（数据表 16、17）：
• pass_selectivity：233 个（Δsel_dock>0 且 Δsel_MMGBSA≥2）
• Tier 1′：57 个（pose + potency + pass_selectivity + contact）
• has_selectivity_contact：63 个（铰链 H-bond 代理）
这些标签不能可靠预测真实的 JNK1/JNK2/JNK3 选择性，不参与 MD 短名单或采购排序。

5.1 Gly87 占据假说
JNK1 在铰链区有甘氨酸 Gly87，JNK2/JNK3 分别为 Ser/Met，体积更大。曾假设：配体占据 JNK1 特有的小空间可得选择性（机制类比参考文献 [10]）。
回顾测试显示：所有 benchmark 配体距 Gly87 仅 0.59–1.18 Å，E1 与 CC-930 均显示占据，无法区分（数据表 08）。策略放弃。

5.2 MM-GBSA 选择性门槛
曾要求 Δsel_MMGBSA ≥ 2 kcal/mol。标定发现：九个对照的 |Δsel| 中位数噪声达 8.13 kcal/mol，2.0 门槛远低于噪声，方向准确率仍仅 43%（数据表 07）。
注意：MMGBSA_JNK1 ≤ −51.6 作为「JNK1 单点活性」门槛仍用于 MD 短名单，与选择性差值无关。
"""),
    ("六、MD 短名单与成药性：如何选出 25 个、再做 16 个 MD？", """
6.1 MD 短名单漏斗（数据表 18）
3125 个通过 pose 质量 → 182 个通过 JNK1 活性双门槛（Glide ≤ −7.43 且 MMGBSA_JNK1 ≤ −51.6）→ 157 个 F1∧F2 → ADMET 剔除 9 个 → 25 个 shortlist。
25 个按化学策略分四组：G1 文献 chemotype 邻近组(9)、G2 新骨架(10)、G3 已知活性对照(4)、G4 阴性锚点(2)。
按组配额取 16 个做 Desmond 分子动力学（48 个任务 = 16 分子 × 3 蛋白）。组内排序键：pose QC → score_JNK1（越低越好）→ MMGBSA_JNK1 单点活性 → 骨架多样性；不用 Δsel/pass_selectivity/Tier（技术报告 §6.1.1）。
G1 进 MD：690、2232、2157、2389；G2 进 MD：**2231、1280、4795、2747、1555、1762**（见 `data/shortlist/md_shortlist_final.csv`）；G1 落选 5 人、G2 落选 4 人完整 ID 待 `candidates_ranked_befe.csv` 归档（数据表 25）。16 人 Δsel_dock 排序与 HIT 报价见数据表 27。

6.2 为什么 2231 能进 MD 但 pass_selectivity=No？
2231 的 JNK1 对接分极强（−11.22），满足活性门槛；未过 pass_selectivity 是因为探索性双门槛中的 Δsel_MMGBSA 标签，MD 短名单从不读取该标签（详见技术报告 §6.3.1）。
"""),
    ("七、分子动力学 QC：结合姿势稳不稳？", """
对每个分子在三个亚型各跑一条 MD 轨迹（约 20–50 ns），用两项指标质检：
• 配体 RMSD ≤ 3 Å：分子在口袋里没有「飞出去」
• 铰链氢键占有率 ≥ 30%：与铰链区有稳定极性接触

pass_md_overall = JNK1 通过 且 (JNK2 或 JNK3 至少一个通过)。
结果：G1 组 3/4 通过，G2 组 0/6 通过——但 G2 的 2231 在 JNK1 上 hinge=0.91、RMSD 最低，仍有探索价值（图 5、数据表 20–21）。

重要：SP600125、E1 等已知活性药在 MD 指标上可 fail，说明 MD 不能代替活性实验，只能辅助判断 pose 可信度。
"""),
    ("八、化合物 2231 的 200 ns 延伸 MD", """
对 MD 相对排名第一的 G2 分子 2231，补做 200 ns Amber 模拟（三 isoform 各一条轨迹）。
与 Desmond 短 MD 对比：JNK1 配体 RMSD 中位数最低（Amber 0.57 Å vs Desmond 0.48 Å），JNK2 最高（图 4，数据表 12）。
JNK1 中 Asn108–配体氧原子氢键占有率约 68%（技术报告 §6.5.4）。

限制：仅单条轨迹、生产期对配体有位置约束、不能与 Desmond 数值直接对比绝对值。
结论：支持「2231 在 JNK1 中 pose 相对更稳」的方向性假说，但不能确认选择性，也不能改写 §6.2 中 overall MD fail 的记录。
"""),
    ("九、最终采购的 10 个分子与实验计划", """
采购清单见数据表 11，分为：
• G3 对照 4 个：SP600125、CC-90001、CC-930、E1——用于校准实验体系
• G1 主力 3 个：690、2232、2157——MD 总体通过、接近文献 chemotype
• G2 探索 3 个：2231、1280、4795——新骨架与 off-target pose 假说

花钱买的不是「已经算出的选择性 hit」，而是：
1) 验证计算管线能否富集有活性分子；2) 比较 G1 与 G2 哪类骨架更易出活性；3) 用 2231/2157 等检验 JNK1 偏好假说。
必做实验：同批次 JNK1 + JNK2 + JNK3 重组酶 IC50。
"""),
    ("十、诚实结论（给决策者与合作者）", """
1. 受体准备与对接流程可信（5/5 再对接）。
2. 机器学习能粗筛家族活性，不能判断亚型方向。
3. Glide Δsel 与 MM-GBSA Δsel 不能用于 isoform 分型（方向准确率 43% 或更低）。
4. MD 可辅助看结合姿势，不能单独确认选择性。
5. 没有任何分子被计算「确认」为 JNK1 选择性抑制剂；选择性只能由湿实验回答。

若只有经费买 2 个分子优先赌 JNK1 选择性：MD 可视化支持 2231 + 2157 的组合；690 更适合作为 pan-JNK 活性验证。
"""),
    ("参考文献与方法的准确对应", """
以下编号与项目技术报告 §11 一致；未列入者为本项目使用的商业软件/通用规则，不强行编造文献。

[1] ChEMBL 数据库 — 训练数据来源。Zdrazil B, et al. Nucleic Acids Res. 2024;52(D1):D1180-D1192. doi:10.1093/nar/gkad1004

[2] CC-90001 临床候选与 JNK1 功能偏向。Bennett BL, et al. J. Med. Chem. 2021;64(3):1776-1795. doi:10.1021/acs.jmedchem.0c01843. PMID:33404223

[3] SP600125 经典 pan-JNK 工具药。Bennett BL, et al. Proc. Natl. Acad. Sci. USA 2001;98(24):13681-13686. doi:10.1073/pnas.251194298. PMID:11717429

[4] E1 与 JNK1 抑制剂系列。Pan X, et al. J. Med. Chem. 2024. doi:10.1021/acs.jmedchem.4c01764

[5] TCS JNK 6O 氨基吡唑 JNK 抑制剂。Szczepankiewicz BG, et al. J. Med. Chem. 2006;49(14):3563-3566. doi:10.1021/jm060150w

[6] CC-930 共晶结构 3TTI。Plantevin-Krenitsky V, et al. Bioorg. Med. Chem. Lett. 2012;22(3):1433-1438. doi:10.1016/j.bmcl.2011.12.111

[7] Glide 分子对接。Friesner RA, et al. J. Med. Chem. 2004;47(7):1739-1749. doi:10.1021/jm0306430

[8] JNK 成药靶点综述。Manning BD, Davis RJ. Nat. Rev. Drug Discov. 2003;2(7):554-565. doi:10.1038/nrd1132

[9] XGBoost 机器学习。Chen T, Guestrin C. Proc. 22nd ACM SIGKDD 2016. doi:10.1145/2939672.2939785

[10] JNK2/3 选择性共价抑制剂与 Gly87/Leu106 机制讨论。Bennett BL, et al. J. Med. Chem. 2022. doi:10.1021/acs.jmedchem.2c01834

方法—工具补充说明（无额外编造文献）：
• MM-GBSA、Desmond MD、QikProp ADMET：Schrödinger 商业软件套件（Prime / Desmond / QikProp），流程参数见项目配置与工作区归档报告。
• 2231 延伸 MD：AMBER 分子动力学 + cpptraj 分析 + MMPBSA.py；力场与输入见 results/md_2231_200ns/README.md。
• Lipinski 五规则、Murcko 骨架划分、Butina 聚类、RDKit：药物化学与化学信息学常用开源/工业界标准流程；本项目按脚本实现，详见仓库 scripts/ 与 data/processed/。
• JNK-IN-8 共价抑制剂对照来源：Cell Chem Biol 2012（见 data/benchmarks/literature_benchmarks.csv 的 source 字段）。
"""),
  ]

  for heading, body in sections:
    doc.add_heading(heading, level=1)
    for para in body.strip().split("\n"):
      if para.strip():
        doc.add_paragraph(para.strip())
    doc.add_paragraph()

  doc.add_heading("附图与附表", level=1)
  add_image(doc, fig_paths["funnel"], "图1 端到端筛选漏斗各阶段化合物数量（数据表 00）")
  add_table(doc, pd.read_csv(DATA_DIR / "00_端到端漏斗汇总.csv"), "表1 端到端漏斗汇总")
  add_image(doc, fig_paths["redock"], "图2 五个蛋白结构的共晶再对接 RMSD（数据表 04）")
  add_table(doc, pd.read_csv(DATA_DIR / "04_再对接验证redocking.csv"), "表2 再对接验证明细")
  add_image(doc, fig_paths["dir_acc"], "图3 Benchmark 方向准确率对比（数据表 19）")
  add_image(doc, fig_paths["bench_dir"], "图4 四个关键对照的 Δsel_dock 与方向匹配（数据表 06）")
  add_table(doc, pd.read_csv(DATA_DIR / "06_benchmark方向混淆矩阵.csv"), "表3 Benchmark 方向混淆")
  add_table(doc, pd.read_csv(DATA_DIR / "22_ML模型性能汇总.csv"), "表4 ML 模型 holdout 性能")
  add_table(doc, pd.read_csv(DATA_DIR / "24_外部decoy验证指标汇总.csv"), "表5 外部 decoy 验证指标")
  add_table(doc, pd.read_csv(DATA_DIR / "18_MD短名单漏斗.csv"), "表6 MD 短名单漏斗")
  add_table(doc, pd.read_csv(DATA_DIR / "26_对接后筛选漏斗.csv"), "表6b 对接后主线筛选漏斗（§3.5）")
  add_table(doc, pd.read_csv(DATA_DIR / "17_Tier分布.csv"), "表7 VSW Tier 分布（选择性探索，未用于采购）")
  add_image(doc, fig_paths["hinge"], "图5 采购分子铰链氢键占有率（数据表 20）")
  add_table(doc, pd.read_csv(DATA_DIR / "11_采购清单purchase_after_md.csv"), "表8 最终采购清单（完整）")
  add_image(doc, fig_paths["2231_rmsd"], "图6 化合物 2231 延伸 MD 配体 RMSD 对比（数据表 12）")
  add_table(doc, pd.read_csv(DATA_DIR / "12_2231延伸MD_RMSD分位数.csv"), "表9 2231 延伸 MD RMSD 统计")
  add_table(doc, pd.read_csv(DATA_DIR / "13_2231延伸MD_MMGBSA分量.csv"), "表10 2231 MM-GBSA 分量（辅助记录，不作选择性裁决）")

  add_table(doc, pd.read_csv(DATA_DIR / "01_文献benchmark化合物.csv"), "表11 九个文献 benchmark 化合物与实验活性")
  add_table(doc, pd.read_csv(DATA_DIR / "08_Gly87自检.csv"), "表12 Gly87 占据策略回顾性测试")
  add_table(doc, pd.read_csv(DATA_DIR / "16_VSW选择性探索各阶段数量.csv"), "表13 VSW 选择性探索各过滤阶段数量")
  add_table(doc, pd.read_csv(DATA_DIR / "20_采购分子铰链占有率.csv"), "表14 采购分子铰链氢键占有率")
  add_table(doc, pd.read_csv(DATA_DIR / "21_采购分子配体RMSD中位数_Angstrom.csv"), "表15 采购分子配体 RMSD 中位数 (Å)")

  doc.add_heading("数据文件索引", level=1)
  doc.add_paragraph(
    f"全部 {len(list(DATA_DIR.glob('*.csv')))} 个 CSV 与 JSON 数据表位于：docs/popular_science/data_tables/\n"
    "详见该文件夹 README.md。"
  )

  doc.save(DOC_PATH)
  print(f"Wrote {DOC_PATH}")


def add_paragraphs(doc: Document, paragraphs: list[str]) -> None:
  for para in paragraphs:
    text = para.strip()
    if text:
      doc.add_paragraph(text)
  doc.add_paragraph()


def add_data_file_to_doc(doc: Document, path: Path, title: str) -> None:
  """Embed one data file in the Word appendix.

  CSV files are inserted as tables. JSON files are inserted as formatted text so
  that nested structures remain readable instead of being flattened incorrectly.
  """
  doc.add_heading(title, level=3)
  if path.suffix.lower() == ".csv":
    df = pd.read_csv(path)
    add_table(doc, df, f"{title}（完整表）")
  elif path.suffix.lower() == ".json":
    obj = json.loads(path.read_text(encoding="utf-8"))
    text = json.dumps(obj, ensure_ascii=False, indent=2)
    for chunk_start in range(0, len(text), 2500):
      p = doc.add_paragraph(text[chunk_start:chunk_start + 2500])
      for r in p.runs:
        r.font.name = "Courier New"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        r.font.size = Pt(8)
  else:
    doc.add_paragraph(path.read_text(encoding="utf-8")[:5000])
  doc.add_paragraph()


def build_word_detailed(fig_paths: dict[str, Path]) -> None:
  """Build a prose-style Word report for non-CADD readers."""
  doc = Document()
  style = doc.styles["Normal"]
  style.font.name = "宋体"
  style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
  style.font.size = Pt(11)

  title = doc.add_heading("JNK1/2/3 抑制剂计算筛选项目通俗说明报告", 0)
  title.alignment = WD_ALIGN_PARAGRAPH.CENTER
  doc.add_paragraph("版本：2.8（成段叙述版，面向完全不懂计算机辅助药物设计的读者）")
  doc.add_paragraph("数据来源：项目仓库 JNK1_Selectivity_Project；全部数值均可回溯到本文件配套的 data_tables/ 文件夹。")
  doc.add_paragraph("技术报告原文：docs/JNK1_PROJECT_REPORT.md（v2.8）")
  doc.add_page_break()

  sections = [
    ("1. 为什么要做这个项目：用计算机先替实验“筛一遍”",
     [
       "本项目研究的是 JNK（c-Jun N-terminal kinase）蛋白家族中的三个亚型：JNK1、JNK2 和 JNK3。可以把这三个亚型理解为同一家族里的三把“锁”，它们的钥匙孔非常相似，但在某些细节上又不完全相同。药物研发希望找到能够插入这些钥匙孔的小分子，其中最理想的情况是分子主要作用于 JNK1，而尽量少作用于 JNK2 和 JNK3。这样做的目的，是希望在保留治疗作用的同时降低不必要的副作用。",
       "直接把成千上万个分子全部买回来做实验，成本高、周期长，而且很多分子很可能根本没有活性。因此本项目采用计算机辅助药物设计（CADD）的思路：先用公开数据库和计算模型，把明显不合适的分子排除掉，再把最值得验证的一小批分子交给湿实验。这里的“湿实验”指真实的生化实验，例如同批次测定 JNK1、JNK2、JNK3 的 IC50。IC50 是评价药物活性的常用指标，数值越小，通常说明抑制作用越强。",
       "需要特别说明的是，本项目最开始的目标是寻找 JNK1 选择性分子，但在计算过程中，我们用文献中已知活性的对照分子反复验证后发现：当前使用的机器学习、分子对接、MM-GBSA 和 MD 指标都不能可靠判断一个分子到底更偏向 JNK1、JNK2 还是 JNK3。因此，最终报告采用了更诚实也更稳妥的表述：计算流程用于富集可能的 JNK 家族结合剂，并帮助判断结合姿势是否可信；真正的亚型选择性只能由同批次 JNK1/JNK2/JNK3 酶学实验回答。",
     ]),
    ("2. 阅读前必须理解的核心名词",
     [
       "“JNK”是一类蛋白激酶。激酶可以把磷酸基团转移到其他蛋白上，从而改变细胞信号。JNK 与炎症、纤维化、细胞应激等过程有关。JNK1、JNK2、JNK3 是这个家族的三个亚型，也就是结构相似但生物学角色不同的蛋白版本。JNK1 的基因名是 MAPK8，JNK2 是 MAPK9，JNK3 是 MAPK10。",
       "“小分子”“配体”“抑制剂”在本报告中经常交替出现。小分子是候选化合物；当它被放进蛋白结合口袋中讨论时，常称为配体；如果它能降低蛋白活性，就称为抑制剂。“pan-JNK”表示一个分子对多个 JNK 亚型都有作用，而不是只对某一个亚型有作用。",
       "“分子对接”可以理解为用计算机把小分子摆进蛋白的结合口袋，并给这个摆放结果打分。Glide XP 是本项目使用的对接方法，参考文献为 Friesner 等在 Journal of Medicinal Chemistry 发表的 Glide 方法论文 [7]。对接分数越负，通常表示计算上预测结合越强，但这个分数不是实验活性，尤其不能直接拿不同蛋白结构之间的分数差来确定选择性。",
       "“机器学习”是让模型从已有分子和活性数据中学习结构与活性的关系。本项目用 XGBoost 建立 JNK1、JNK2、JNK3 三个模型。XGBoost 的方法来源是 Chen 和 Guestrin 2016 年 KDD 论文 [9]。机器学习在这里主要用于粗筛 JNK 家族活性，而不是判断精细的亚型选择性。",
       "“MD”是分子动力学模拟。它不是简单给分子打分，而是让蛋白和小分子在计算机里随时间运动，观察小分子是否能在口袋里保持稳定。RMSD 是 MD 中常用的指标，表示结构相对初始位置偏离多少，单位是埃（Å）。铰链氢键占有率表示小分子在模拟过程中有多少比例的时间与激酶铰链区保持氢键接触。",
       "“ADMET”是药物研发中对吸收、分布、代谢、排泄和毒性的总称。本项目用 QikProp 预测 hERG、口服吸收、Caco-2 通透性、溶解度等指标。QikProp、Desmond、Prime MM-GBSA 属于 Schrödinger 商业软件套件；本报告只说明其用途和项目参数，不为商业软件额外编造文献。",
     ]),
    ("3. 总体流程：从 5000 个分子到 10 个采购",
     [
       "整个筛选流程可以概括为一条主线：机器学习粗筛 → Glide XP 三亚型对接 → pose QC 与 JNK1 活性筛选 → ADMET 成药性过滤 → 分子动力学验证 → 采购 10 个分子做实验。端到端数量是：4983 个进入对接库，4979 个有有效对接结果，157 个进入 ADMET 前短名单，25 个进入 ADMET 后 shortlist，16 个进入 MD，最终推荐 10 个采购。",
       "主线筛选不使用 pass_selectivity、Tier 1′、Tier 2 或 Δsel 方向，而是用 JNK1 活性门槛、pose 质量、ADMET 和 MD pose QC 来选择要做实验的分子。这解释了为什么 2231 虽然是 Tier 3、pass_selectivity 为 No，却仍然可以进入 MD：2231 凭 JNK1 对接分和 pose QC 满足 MD 短名单条件，而非依赖选择性探索标签。",
       "项目早期曾尝试用对接分、MM-GBSA 差值和接触特征给分子打 pass_selectivity 或 Tier 标签，得到 233 个 pass_selectivity 分子、Tier 1′ 为 57 个。但后续 benchmark 验证证明这些标签不能可靠预测真实的 JNK1/JNK2/JNK3 选择性。这些失败的探索尝试完整保留于技术报告第五章，供回顾分析，但不参与采购决策。",
     ]),
    ("4. 第一步：机器学习粗筛，目标是“不漏掉可能有活性的分子”",
     [
       "机器学习模型的数据来自 ChEMBL 数据库 [1]。ChEMBL 是公开的药物化学数据库，收录了大量化合物及其对靶点的活性数据。本项目清洗后得到 JNK1 444 个化合物、JNK2 610 个化合物、JNK3 1147 个化合物，并分别训练三个 XGBoost 回归模型。模型预测的是 pActivity，也就是活性的对数尺度；数值越高，通常表示活性越强。",
       "为了避免模型只记住相似分子而表现虚高，数据划分采用 Murcko 骨架分组。简单说，就是把化学骨架相似的分子尽量放在同一数据子集中，避免训练集中出现一个分子、测试集中出现几乎同样的分子。最终 holdout 测试结果为：JNK1 R²=0.697，JNK2 R²=0.574，JNK3 R²=0.774；Spearman 相关分别为 0.858、0.780、0.869。这些数值说明模型对“是否可能有家族活性”有一定帮助。",
       "模型第一道门槛是 p_family ≥ 6.0。这个门槛的设计不是为了精准判断谁一定有活性，而是为了尽量不漏掉已知活性分子。9 个文献 benchmark 分子在这个阈值下全部通过，说明召回能力好。但后来我们补充 10,000 个 Taosu decoy 做外部验证，发现 decoy 假阳性率达到 95.3%，specificity 只有 4.7%。这说明 F1 是一张非常宽的网：好处是不会轻易漏掉真活性，缺点是会放进大量假阳性。",
       "因此，机器学习在本项目中的定位是粗筛和排序，不是最终裁决。真正缩小范围还要依赖 final_score、对接、ADMET 和 MD。final_score 把 p_family、JNK1 预测活性、QED 和合成可及性 SA 组合起来排序，所以 Top-5000 不是单纯按 p_family 取前 5000。",
       "为什么机器学习不能直接判断 JNK1 选择性？原因是可用于训练选择性分类器的数据太少。项目数据中 JNK1-selective 标注只有 8 个，选择性分类器测试 F1=0。对 E1 和 TCS JNK 6O 这些关键文献对照，ML 预测最高亚型也与实验不一致。因此报告明确写出：ML 不用于 isoform 方向判断。",
     ]),
    ("5. 第二步：Glide XP 分子对接，目标是看分子能否合理放入蛋白口袋",
     [
       "对接使用每个 JNK 亚型一个主结构：JNK1 使用 3ELJ，JNK2 使用 3E7O，JNK3 使用 3TTI。选择这些结构前，项目做了共晶再对接验证。所谓再对接，是把原来在晶体结构里已经看到的配体拿出来，再用计算方法放回同一个蛋白口袋，看能否复现实验姿势。若重对接 pose 与晶体 pose 的 RMSD 小于 2 Å，通常说明结构准备和网格设置是可信的。",
       "本项目 5 个结构再对接全部通过：3ELJ 为 0.66 Å，4L7F 为 0.92 Å，3E7O 为 0.26 Å，3TTI 为 1.50 Å，4WHZ 为 1.88 Å。这个结果支持用这些 PDB 建立对接网格。但必须强调，再对接通过只说明“能复现已知共晶 pose”，并不说明对接分数可以区分 JNK1/JNK2/JNK3 选择性。",
       "对接后得到 4979 个有效记录。主线采购链的逐层筛选为（数据表 26，技术报告 §3.5）：pass_pose 3234（VSW pose 质量门）→ MD-F1 3125（更严 pose QC）→ MD-F2 182（score_JNK1 ≤ −7.43 且 MMGBSA_JNK1 ≤ −51.6）→ F1∧F2 共 157 个。pass_potency 1681 为单 Glide 活性门，可与 pose 独立计数。",
       "早期项目定义了 Δsel_dock = min(score_JNK2, score_JNK3) − score_JNK1。因为 Glide 分越负通常表示结合越强，所以如果 Δsel_dock 大于 0，计算上看起来像 JNK1 比 JNK2/JNK3 更有优势。这个定义在数学上容易理解，但它有一个根本问题：三个亚型来自不同 PDB、不同网格和不同蛋白环境，绝对对接分并不天然可比。",
       "为了检验 Δsel 是否可信，项目用文献中有 IC50 的 benchmark 分子做回顾验证。结果显示，VSW 单 PDB 口径的方向准确率只有 43%（3/7），归档 ensemble 口径只有 29%（2/7），都低于预设 55% 阈值。也就是说，尽管 Spearman 秩相关看起来有一定数值，真正判断“方向”时仍然不可靠。因此，所有基于 Δsel 的选择性标签都从采购决策中移除。",
     ]),
    ("6. 第三步：为什么 MM-GBSA 和 Gly87 策略也没有成为选择性硬门槛",
     [
       "MM-GBSA 是一种在对接 pose 基础上估算结合自由能的方法。本项目中 MM-GBSA 有两种用途，必须分开理解。第一种是 MMGBSA_JNK1 单点活性门槛，用于判断一个分子是否可能在 JNK1 上有足够结合能力，门槛是 ≤ −51.6。第二种是 Δsel_MMGBSA，也就是比较 JNK1 与 JNK2/JNK3 的差值，早期曾用于 pass_selectivity。",
       "Benchmark 标定显示，Δsel_MMGBSA 的 |Δsel| 中位数噪声为 8.13 kcal/mol，而早期门槛只有 2.0 kcal/mol，远低于噪声水平。MM-GBSA 方向准确率也只有 43%。因此，Δsel_MMGBSA 不再作为选择性硬筛。但 MMGBSA_JNK1 单点活性门槛仍可用于 MD 短名单，因为它回答的是“这个分子是否可能结合 JNK1”，而不是“它是否选择性偏向 JNK1”。",
       "项目还尝试过 Gly87 占据策略。JNK1 在铰链附近有 Gly87，JNK2/JNK3 对应位置是 Ser/Met，体积和化学性质不同。理论上，如果一个分子占据 JNK1 特有的小空间，可能带来 JNK1 选择性。这个思路与文献中 JNK2/3 选择性机制的讨论有关 [10]。但回顾测试发现，E1、TCS JNK 6O、CC-930、SP600125、CC-90001 等 benchmark 都靠近 Gly87，距离约 0.59–1.18 Å，无法区分 JNK1 选择性与 pan 或 JNK2/3 偏好。因此 Gly87 指标被放弃。",
     ]),
    ("7. 第四步：MD 短名单和 ADMET，目标是把候选变成可以买、可测的分子",
     [
       "MD 短名单不是从 pass_selectivity 来的，而是从实际可用于采购和验证的角度设计。首先，4983 个 F0 分子经过 pose QC，3125 个通过；再经过 JNK1 活性和配体效率过滤，182 个通过；F1 与 F2 同时满足后为 157 个；再经过 QikProp ADMET 过滤和分组策略，得到 25 个 shortlist。",
       "ADMET 过滤的原因很实际：即使一个分子在对接中得分很好，如果预测溶解性很差、hERG 风险高、口服吸收差，后续实验和成药开发价值也会降低。本项目考虑 hERG、口服吸收、Caco-2 通透性、溶解度和 #stars 等指标。G3 文献对照分子即使 ADMET 不完美也会保留，因为它们的作用是校准实验体系，而不是作为新药候选。",
       "25 个 shortlist 被分成四组：G1 是相对接近文献 chemotype 的分子，G2 是新骨架，G3 是已知活性对照，G4 是阴性锚点。进入 MD 的 16 个分子按组配额选择：G1 取 4/9，G2 取 6/10，G3 和 G4 全部进入。组内排序以 pose QC、score_JNK1、MMGBSA_JNK1 单点活性和骨架多样性为主，不使用 Δsel 或 pass_selectivity（§6.1.1）。G1 进 MD 的 4 人为 690、2232、2157、2389；G2 进 MD 的 6 人为 2231、1280、4795、2747、1555、1762（数据表 25、`data/shortlist/md_shortlist_final.csv`）。G1 落选 5 人、G2 落选 4 人的完整 ID 仍待 candidates_ranked_befe.csv 归档。16 人按 Δsel_dock 的计算选择性排序与 HIT 报价对照见数据表 27。",
     ]),
    ("8. 第五步：分子动力学 MD，目标是看结合姿势是否稳定",
     [
       "分子对接给出的是一个静态姿势，而真实蛋白和分子在水溶液中会不断运动。因此项目对 16 个分子分别在 JNK1、JNK2、JNK3 中做 MD，共 48 个 Desmond 任务。MD 的核心问题是：这个分子在模拟过程中是否能稳定留在口袋中，是否保持与激酶铰链区的关键接触。",
       "本项目的 MD QC 使用两个主要指标。第一是配体 RMSD，门槛为 ≤ 3 Å；如果 RMSD 很大，通常表示分子姿势漂移甚至离开结合口袋。第二是 hinge H-bond occupancy，门槛为 ≥ 30%；它表示模拟过程中有多少比例的时间保持铰链氢键。pass_md_JNK1 要求 JNK1 上同时满足 RMSD 和 hinge；pass_md_overall 要求 JNK1 pass，并且 JNK2 或 JNK3 至少一个也 pass。",
       "结果显示，G1 组 3/4 通过 overall，G2 组 0/6 通过 overall，G3 对照 2/4 通过，G4 阴性锚点 0/2 通过。这个结果支持 G1 作为更稳妥的活性验证组，也说明 G2 新骨架风险较高。但 MD 不能等同于活性实验，因为 SP600125 和 CC-930 等已知活性分子也可能在 hinge 指标上表现不好。也就是说，MD 是 pose 可信度工具，不是选择性或活性的最终裁判。",
       "在六个新采购分子中，2231 的 JNK1 hinge occupancy 为 0.91、JNK2 为 0.00、JNK3 为 0.10，JNK1 配体 RMSD 也处于最低档。2157 的 JNK1 hinge 为 0.85，JNK3 几乎无 hinge。这些结果使 2231 和 2157 成为优先检验 JNK1 偏好假说的分子；而 690 虽然 Tier 1′ 且 MD overall 通过，但三亚型 hinge 都不低，更像 pan-JNK 结合模式。",
     ]),
    ("9. 2231 的 200 ns Amber 延伸 MD：为什么做、看到了什么、不能说明什么",
     [
       "2231 是一个特殊案例。它在选择性探索中是 Tier 3、pass_selectivity 为 No，但通过 JNK1 活性门槛和 pose QC 进入 25 个 shortlist，并因 G2 配额进入 MD。短 MD 显示它在 JNK1 上 pose 相对更稳，所以项目又为 2231 补做了 200 ns Amber 延伸 MD，三个亚型各一条轨迹。",
       "延伸 MD 的目的不是证明 2231 已经是 JNK1 选择性抑制剂，而是检验短 MD 中“JNK1 pose 更稳”的方向是否可重复观察。50–200 ns 生产期统计显示，2231 在 JNK1 中的配体 heavy-atom RMSD 中位数为 0.57 Å，JNK2 为 1.74 Å，JNK3 为 1.08 Å。这个排序与短 Desmond MD 中 JNK1 最稳、JNK2 相对较差的方向一致。",
       "此外，cpptraj 氢键分析显示，JNK1 中 MOL@O2 与 Asn108@ND2 的接触占有率为 68.4%，是本轨迹中最突出的结构特征。MM-GBSA 分量中 JNK1 ΔG_total 为 −31.9 kcal/mol，JNK2 为 −16.8，JNK3 为 −16.2。但报告明确禁止把这些 MM-GBSA 差值解读为“2231 对 JNK1 有 −15 kcal/mol 选择性优势”，因为 benchmark 已经证明跨亚型 MM-GBSA 差值噪声大，且本次计算有 internal potential 警告。",
       "2231 延伸 MD 的限制也必须写清楚：它仍然是单副本，不足以做统计显著性判断；生产期对配体施加了位置 restraint，因此配体 RMSD 反映的是受约束姿态相对初帧的偏离；Amber 与 Desmond 的力场和分析口径不同，数值不能直接等同。正确结论是：2231 是值得实验验证的 JNK1 偏好假说分子，但计算不能确认其选择性。",
     ]),
    ("10. 最终为什么推荐购买这 10 个分子",
     [
       "最终采购清单包含 10 个分子，并不是因为这 10 个都被计算确认有 JNK1 选择性，而是因为它们共同组成了一个可解释、可验证的实验面板。G3 对照包括 SP600125、CC-90001、CC-930 和 E1，用于确认实验体系能否重现已知活性趋势。没有这些对照，即使新分子测出结果，也很难判断是管线有效还是实验体系本身有偏差。",
       "G1 主力分子包括 690、2232、2157。它们在 MD 中整体表现更稳，且相对接近文献 chemotype。690 同时出现在 Tier 1′、top selective 聚类代表、FEP+ 推荐清单和 panJNK_JNK1bias 子集中，但由于三亚型 hinge 均不低，更适合被描述为 pan-JNK 活性验证分子。2157 虽然 Δsel_dock 为负，但 MD hinge 不对称性支持其作为 JNK1 偏好假说的第二候选。",
       "G2 探索分子包括 2231、1280、4795。G2 overall MD 通过率为 0/6，风险更高，但它们提供了新骨架和 off-target pose 假说。2231 尤其值得关注，因为短 MD 与 200 ns 延伸 MD 都显示 JNK1 pose 相对更稳。1280 和 4795 更像 JNK2/JNK3 或 off-target pose 备份，买它们的目的不是“押中 JNK1”，而是帮助检验计算模型的边界。",
       "因此，花钱购买这 10 个分子的核心逻辑是：用 G3 校准实验，用 G1 检验相对稳妥的家族活性，用 G2 检验新骨架假说，用 G4 阴性锚点检查假阳性。最终真正有无 JNK1 选择性，必须由同批次 JNK1/JNK2/JNK3 IC50 回答。",
     ]),
    ("11. 本项目最重要的负面结论：计算选择性失败本身也是结果",
     [
       "对于完全不熟悉 CADD 的读者，可能会觉得“既然计算不能确认选择性，那是不是项目失败了”。更准确的理解是：计算筛选的价值不只在于给出阳性结论，也在于帮助我们提前识别哪些看似合理的指标其实不能用。本项目用文献 benchmark 否定了 ML 亚型方向预测、Glide Δsel 选择性、MM-GBSA 差值门槛和 Gly87 占据策略。这些负面结论减少了后续实验中的错误期待。",
       "报告最终没有声称“发现了 JNK1 选择性 hit”，而是谨慎地说：得到了一批有明确计算证据和实验验证价值的 JNK 家族候选分子。这样的表述更符合数据事实，也更容易经受审稿或答辩质疑。",
     ]),
    ("12. 参考文献与每种方法的对应关系",
     [
       "ChEMBL 数据库用于训练数据来源，对应参考文献 [1]。CC-90001、SP600125、E1、TCS JNK 6O、CC-930 等对照分子的文献来源分别对应 [2]–[6]。Glide 分子对接方法对应 [7]。JNK 靶点背景与成药挑战对应综述 [8]。XGBoost 机器学习对应 [9]。Gly87/Leu106 选择性机制讨论对应 [10]。",
       "本报告没有为未在技术报告中登记的内容编造文献。Schrödinger Prime MM-GBSA、Desmond、QikProp 被标注为商业软件套件；AMBER/cpptraj/MMPBSA.py 用于 2231 延伸 MD 的方法说明来自项目归档文件。若后续投稿或正式汇报需要补充软件引用，应以实际软件版本和官方引用格式为准，而不是在本报告中临时编造。",
       "[1] Zdrazil B, et al. The ChEMBL Database in 2023. Nucleic Acids Research. 2024;52(D1):D1180-D1192. doi:10.1093/nar/gkad1004。用途：ChEMBL 活性数据来源。",
       "[2] Bennett BL, et al. CC-90001. Journal of Medicinal Chemistry. 2021;64(3):1776-1795. doi:10.1021/acs.jmedchem.0c01843. PMID:33404223。用途：CC-90001 背景与功能偏向说明。",
       "[3] Bennett BL, et al. SP600125. Proceedings of the National Academy of Sciences of the USA. 2001;98(24):13681-13686. doi:10.1073/pnas.251194298. PMID:11717429。用途：经典 pan-JNK 工具药对照。",
       "[4] Pan X, et al. Discovery of JNK1 inhibitors for IPF, compound E1. Journal of Medicinal Chemistry. 2024. doi:10.1021/acs.jmedchem.4c01764。用途：E1、Q63 等 JNK1 抑制剂系列与文献对照。",
       "[5] Szczepankiewicz BG, et al. Discovery of aminopyrazole inhibitors of JNK, including TCS JNK 6o. Journal of Medicinal Chemistry. 2006;49(14):3563-3566. doi:10.1021/jm060150w。用途：TCS JNK 6O 文献对照。",
       "[6] Plantevin-Krenitsky V, et al. Discovery of tanzisertib / CC-930 and co-crystal structure 3TTI. Bioorganic & Medicinal Chemistry Letters. 2012;22(3):1433-1438. doi:10.1016/j.bmcl.2011.12.111。用途：CC-930 与 JNK3 共晶结构背景。",
       "[7] Friesner RA, et al. Glide: a new approach for rapid, accurate docking and scoring. Journal of Medicinal Chemistry. 2004;47(7):1739-1749. doi:10.1021/jm0306430。用途：Glide XP 分子对接方法。",
       "[8] Manning BD, Davis RJ. Targeting JNK for therapeutic benefit. Nature Reviews Drug Discovery. 2003;2(7):554-565. doi:10.1038/nrd1132。用途：JNK 作为药物靶点的生物学背景和成药挑战。",
       "[9] Chen T, Guestrin C. XGBoost: A Scalable Tree Boosting System. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining. 2016. doi:10.1145/2939672.2939785。用途：XGBoost 机器学习算法。",
       "[10] Bennett BL, et al. JNK2/3-selective covalent inhibitor YL5084. Journal of Medicinal Chemistry. 2022. doi:10.1021/acs.jmedchem.2c01834。用途：JNK2/3 选择性机制和 Gly87/Leu106 策略讨论。",
     ]),
  ]

  for heading, paragraphs in sections:
    doc.add_heading(heading, level=1)
    add_paragraphs(doc, paragraphs)

  doc.add_heading("13. 图表与核心数据", level=1)
  doc.add_paragraph("以下图表均由 data_tables/ 中的数据自动生成或直接读取。为方便非专业读者阅读，正文先给出核心表；完整数据文件附在下一章。")
  add_image(doc, fig_paths["funnel"], "图1 端到端筛选漏斗各阶段化合物数量（数据表 00）")
  add_table(doc, pd.read_csv(DATA_DIR / "00_端到端漏斗汇总.csv"), "表1 端到端漏斗汇总")
  add_image(doc, fig_paths["redock"], "图2 五个蛋白结构的共晶再对接 RMSD（数据表 04）")
  add_table(doc, pd.read_csv(DATA_DIR / "04_再对接验证redocking.csv"), "表2 再对接验证明细")
  add_image(doc, fig_paths["dir_acc"], "图3 Benchmark 方向准确率对比（数据表 19）")
  add_image(doc, fig_paths["bench_dir"], "图4 四个关键对照的 Δsel_dock 与方向匹配（数据表 06）")
  add_table(doc, pd.read_csv(DATA_DIR / "06_benchmark方向混淆矩阵.csv"), "表3 Benchmark 方向混淆")
  add_table(doc, pd.read_csv(DATA_DIR / "22_ML模型性能汇总.csv"), "表4 ML 模型 holdout 性能")
  add_table(doc, pd.read_csv(DATA_DIR / "24_外部decoy验证指标汇总.csv"), "表5 外部 decoy 验证指标")
  add_table(doc, pd.read_csv(DATA_DIR / "18_MD短名单漏斗.csv"), "表6 MD 短名单漏斗")
  add_table(doc, pd.read_csv(DATA_DIR / "26_对接后筛选漏斗.csv"), "表6b 对接后主线筛选漏斗（§3.5）")
  add_table(doc, pd.read_csv(DATA_DIR / "17_Tier分布.csv"), "表7 VSW Tier 分布（选择性探索，未用于采购）")
  add_image(doc, fig_paths["hinge"], "图5 采购分子铰链氢键占有率（数据表 20）")
  add_table(doc, pd.read_csv(DATA_DIR / "11_采购清单purchase_after_md.csv"), "表8 最终采购清单（完整）")
  add_image(doc, fig_paths["2231_rmsd"], "图6 化合物 2231 延伸 MD 配体 RMSD 对比（数据表 12）")
  add_table(doc, pd.read_csv(DATA_DIR / "12_2231延伸MD_RMSD分位数.csv"), "表9 2231 延伸 MD RMSD 统计")
  add_table(doc, pd.read_csv(DATA_DIR / "13_2231延伸MD_MMGBSA分量.csv"), "表10 2231 MM-GBSA 分量（辅助记录，不作选择性裁决）")
  add_table(doc, pd.read_csv(DATA_DIR / "01_文献benchmark化合物.csv"), "表11 九个文献 benchmark 化合物与实验活性")
  add_table(doc, pd.read_csv(DATA_DIR / "08_Gly87自检.csv"), "表12 Gly87 占据策略回顾性测试")
  add_table(doc, pd.read_csv(DATA_DIR / "16_VSW选择性探索各阶段数量.csv"), "表13 VSW 选择性探索各过滤阶段数量")
  add_table(doc, pd.read_csv(DATA_DIR / "20_采购分子铰链占有率.csv"), "表14 采购分子铰链氢键占有率")
  add_table(doc, pd.read_csv(DATA_DIR / "21_采购分子配体RMSD中位数_Angstrom.csv"), "表15 采购分子配体 RMSD 中位数 (Å)")

  doc.add_heading("14. 完整数据附录（00–26 号文件逐一嵌入）", level=1)
  doc.add_paragraph("本章把 data_tables/ 文件夹中的 00–26 号 CSV/JSON 文件逐一嵌入 Word。若表格较宽，建议在 Word 中横向页面查看；原始机器可读文件同时保存在 data_tables/ 文件夹中。")
  for path in sorted(DATA_DIR.glob("*")):
    if path.name == "README.md" or path.suffix.lower() not in {".csv", ".json"}:
      continue
    add_data_file_to_doc(doc, path, path.name)

  doc.add_heading("15. 数据文件夹说明", level=1)
  doc.add_paragraph("所有数据表格文件位于 docs/popular_science/data_tables/。该文件夹包含 28 个 CSV/JSON 文件和 README.md，既可以供 Word 阅读，也可以供后续 Excel、Python 或统计软件复核。")
  doc.add_paragraph("生成脚本：scripts/build_popular_science_doc.py。重新生成命令：python3 scripts/build_popular_science_doc.py。")

  doc.save(DOC_PATH)
  print(f"Wrote detailed {DOC_PATH}")


def main():
  import argparse
  parser = argparse.ArgumentParser(description="Build popular-science Word doc and figures")
  parser.add_argument("--figures-only", action="store_true", help="Only refresh data_tables and figures")
  args = parser.parse_args()

  setup_data_folder()
  figs = make_figures()
  if not args.figures_only:
    build_word_detailed(figs)
    print(f"Wrote detailed {DOC_PATH}")
  print(f"Data tables: {DATA_DIR}")
  print(f"Figures: {FIG_DIR}")


if __name__ == "__main__":
  main()
